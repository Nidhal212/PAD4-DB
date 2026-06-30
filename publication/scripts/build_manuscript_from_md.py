"""
build_manuscript_from_md.py — Render PAD4_DB_manuscript_FINAL.md to .docx with figures.

Detects:
  - Pipe tables                        → styled Word table
  - **Figure N | caption**             → inserts PNG then styled caption
  - **Supplementary Figure SN | cap** → inserts PNG then styled caption
  - Headings (# / ## / ###)            → heading paragraphs
  - Numbered list (references)         → hanging-indent 9pt
  - Bullet lists                       → bullet 10pt
  - Horizontal rules (---)             → skipped
  - Bold/italic/code inline            → inline formatting

Run: conda run -n pad4bench python publication/scripts/build_manuscript_from_md.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path('/home/nidhal/PAD4-db_V2')
MD   = ROOT / 'publication/manuscript/PAD4_DB_manuscript_FINAL.md'
OUT  = ROOT / 'publication/manuscript/PAD4_DB_manuscript_FINAL.docx'

FIGS = ROOT / 'publication/figures'

# Map figure label → path (relative to FIGS dir, no extension)
FIG_MAP = {
    'Figure 1':                  'main/fig01_headline',
    'Figure 2':                  'main/fig02_source_overlap',
    'Figure 3':                  'main/fig03_potency',
    'Figure 4':                  'main/fig04_scaffold',
    'Figure 5':                  'main/fig05_cliff_network',
    'Figure 6':                  'main/fig06b_cliff_pairs',
    'Supplementary Figure S1':   'supplementary/fig_s01_scaffold_cliff_density',
    'Supplementary Figure S2':   'supplementary/fig_s02_assay_enrichment',
    'Supplementary Figure S3':   'supplementary/fig_s03_sas_map',
    'Supplementary Figure S4':   'supplementary/fig_s04_ruggedness_panels',
    'Supplementary Figure S5':   'supplementary/fig_s05_null_models',
}

# Max content width: 21cm page - 2×2.4cm margins = 16.2cm; use 15.8cm
MAX_WIDTH_CM = 15.8

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(2.4)
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)


# ── Inline formatting helpers ──────────────────────────────────────────────────

_TOKEN = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')

def _set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_inline(p, text, size=11, base_italic=False):
    """Add text to paragraph with **bold**, *italic*, `code` inline markup."""
    for seg in _TOKEN.split(text):
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            _set_font(p.add_run(seg[2:-2]), size=size, bold=True, italic=base_italic)
        elif seg.startswith('*') and seg.endswith('*'):
            _set_font(p.add_run(seg[1:-1]), size=size, italic=True)
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1])
            _set_font(r, size=size - 0.5, italic=base_italic)
            r.font.name = 'Consolas'
        else:
            _set_font(p.add_run(seg), size=size, italic=base_italic)


# ── Block-level helpers ────────────────────────────────────────────────────────

def para(text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6,
         indent=None, hanging=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
        if hanging:
            p.paragraph_format.first_line_indent = Cm(-indent)
    add_inline(p, text, size=size, base_italic=italic)
    return p


def heading(text, level):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    _set_font(p.add_run(text), size=sizes.get(level, 11),
              bold=True, italic=(level == 3))
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after  = Pt(4)


# ── Table helpers ──────────────────────────────────────────────────────────────

def _set_tbl_borders(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom'):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '8')
        e.set(qn('w:color'), '000000')
        borders.append(e)
    ih = OxmlElement('w:insideH')
    ih.set(qn('w:val'), 'single')
    ih.set(qn('w:sz'), '2')
    ih.set(qn('w:color'), 'CCCCCC')
    borders.append(ih)
    tblPr.append(borders)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), hex_color)
    sh.set(qn('w:val'), 'clear')
    tcPr.append(sh)


def add_table(rows):
    """Render a list of string-lists as a styled Word table."""
    if not rows:
        return
    ncol = len(rows[0])
    tbl  = doc.add_table(rows=len(rows), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_tbl_borders(tbl)
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            val = row[j] if j < len(row) else ''
            add_inline(cell.paragraphs[0], val, size=8.5)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8.5)
            if i == 0:   # header row: dark blue bg, white text
                _shade_cell(cell, '1A237E')
                for r in cell.paragraphs[0].runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:  # alternate rows: light blue tint
                _shade_cell(cell, 'EEF2FF')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Figure helpers ─────────────────────────────────────────────────────────────

def insert_figure(fig_key, caption_text):
    """Insert figure PNG then caption text."""
    rel = FIG_MAP.get(fig_key)
    if not rel:
        para(f'[FIGURE NOT FOUND: {fig_key}]', size=9, italic=True)
        para(caption_text, size=9, italic=True, after=12)
        return

    png = (FIGS / rel).with_suffix('.png')
    if not png.exists():
        para(f'[FILE NOT FOUND: {png.name}]', size=9, italic=True)
        para(caption_text, size=9, italic=True, after=12)
        return

    # Compute display width: fit within MAX_WIDTH_CM, preserve aspect ratio
    w, h = Image.open(png).size
    aspect = w / h
    if aspect >= 1:   # landscape → constrain width
        width_cm = MAX_WIDTH_CM
    else:             # portrait → constrain height to ~20cm, derive width
        width_cm = min(MAX_WIDTH_CM, 20.0 * aspect)

    # Space before figure
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(10)
    sp.paragraph_format.space_after  = Pt(4)

    # Insert image centred
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(2)
    img_p.paragraph_format.space_after  = Pt(4)
    img_p.add_run().add_picture(str(png), width=Cm(width_cm))

    # Caption paragraph (bold label + italic rest)
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_after  = Pt(14)
    cap_p.paragraph_format.space_before = Pt(2)

    # Split label ("Figure N | ") from body text
    label_re = re.match(
        r'^(\*\*(?:Supplementary )?Figure S?\d+\s*\|[^*]*\*\*)\s*(.*)',
        caption_text, re.DOTALL
    )
    if label_re:
        label_md = label_re.group(1)
        body_md  = label_re.group(2)
        # render bold label
        label_txt = label_md.strip('*')
        _set_font(cap_p.add_run(label_txt), size=9, bold=True)
        # render body as italic
        if body_md.strip():
            _set_font(cap_p.add_run(' ' + body_md.strip()), size=9, italic=True)
    else:
        add_inline(cap_p, caption_text, size=9, base_italic=True)


# ── Regex patterns ─────────────────────────────────────────────────────────────

# Matches: **Figure 1 | ...** or **Supplementary Figure S3 | ...**
FIG_LINE_RE = re.compile(
    r'^\*\*((?:Supplementary )?Figure S?\d+)\s*\|'
)

# Matches separator-like: --- or ===
HR_RE = re.compile(r'^-{3,}$|^={3,}$')


# ── Parse and render ───────────────────────────────────────────────────────────

lines = MD.read_text(encoding='utf-8').split('\n')
i, first_h1 = 0, True

while i < len(lines):
    ln = lines[i].rstrip('\n')
    st = ln.strip()

    # blank line or horizontal rule → skip
    if st == '' or HR_RE.match(st):
        i += 1
        continue

    # pipe table: collect consecutive pipe rows
    if st.startswith('|'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i].strip())
            i += 1
        rows = []
        for r in block:
            if re.match(r'^\|[\s:|-]+\|$', r):   # separator row
                continue
            # Replace escaped pipes \| with placeholder before splitting,
            # then restore so cells contain literal | characters.
            r_safe = r.replace('\\|', '\x00PIPE\x00')
            cells = [c.strip().replace('\x00PIPE\x00', '|')
                     for c in r_safe.strip('|').split('|')]
            rows.append(cells)
        if rows:
            add_table(rows)
        continue

    # figure caption line: **Figure N | caption text**
    fig_match = FIG_LINE_RE.match(st)
    if fig_match:
        fig_key = fig_match.group(1)   # e.g. "Figure 1" or "Supplementary Figure S3"
        insert_figure(fig_key, st)
        i += 1
        continue

    # legacy blockquote captions (kept for safety)
    if st.startswith('>'):
        body = st.lstrip('> ').strip()
        bm = re.search(r'\[((?:Supplementary )?Figure S?\d+)[a-z]?\.', body)
        if bm:
            insert_figure(bm.group(1), body)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_inline(p, body, size=9, base_italic=True)
        i += 1
        continue

    # headings
    if st.startswith('### '):
        heading(st[4:], 3); i += 1; continue
    if st.startswith('## '):
        heading(st[3:], 2); i += 1; continue
    if st.startswith('# '):
        if first_h1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_font(p.add_run(st[2:]), size=16, bold=True)
            first_h1 = False
        else:
            heading(st[2:], 1)
        i += 1; continue

    # numbered list (references section)
    if re.match(r'^\d+\.\s', st):
        body = re.sub(r'^\d+\.\s', '', st)
        num  = st.split('.', 1)[0] + '. '
        p    = para(body, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
                    after=3, indent=0.6, hanging=True)
        # prepend number to first run
        p.runs[0].text = num + p.runs[0].text
        i += 1; continue

    # bullet list
    if st.startswith('- ') or st.startswith('* '):
        para('•  ' + st[2:], size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
             after=3, indent=0.5, hanging=True)
        i += 1; continue

    # plain paragraph
    para(st)
    i += 1


doc.save(str(OUT))

# Stats
import zipfile
n_img = len([n for n in zipfile.ZipFile(OUT).namelist()
             if n.startswith('word/media/')])
size_kb = OUT.stat().st_size / 1024
print(f"Saved {OUT.name}: {size_kb:.0f} KB · {len(doc.tables)} tables · {n_img} images")
