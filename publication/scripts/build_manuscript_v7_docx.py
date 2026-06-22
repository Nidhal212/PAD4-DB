"""
build_manuscript_v7_docx.py — Builds the v7 top-journal draft as a Word .docx.

Renders PAD4_DB_manuscript_DRAFT_v7.md (the polished Nature/Cell-style draft) with:
  • Full text: Abstract, Introduction, Results, Discussion, Methods, Conclusion
  • Main figures Fig 1–6 embedded inline at first citation (regenerated v7 figures)
  • Supplementary Fig S10 (SAS map) + S9 (assay enrichment)
  • Tables 1–4 as native Word tables (verified numbers; MMP = 80/85.1%)

Reuses the docx helper conventions from build_manuscript_docx.py.

Run from project root:
  conda run -n pad4bench python publication/scripts/build_manuscript_v7_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path('/home/nidhal/PAD4-db_V2')
FIGS_MAIN = ROOT / 'publication/figures/main'
FIGS_SUPP = ROOT / 'publication/figures/supplementary'
OUT = ROOT / 'publication/manuscript/PAD4_DB_manuscript_DRAFT_v7.docx'

doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.4)
section.top_margin  = section.bottom_margin = Cm(2.4)
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Arial'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 13, 2: 11.5, 3: 11}
    set_run_font(run, size=sizes.get(level, 11), bold=True, italic=(level == 3))
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_md(text, justify=True, size=11, italic=False, space_after=6):
    """Body paragraph with inline **bold** support."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    for i, seg in enumerate(text.split('**')):
        if seg == '':
            continue
        run = p.add_run(seg)
        set_run_font(run, size=size, bold=(i % 2 == 1), italic=italic)
    return p


def add_caption(label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(10)
    r1 = p.add_run(label + ' '); set_run_font(r1, size=9, bold=True)
    r2 = p.add_run(body);        set_run_font(r2, size=9)
    return p


def add_figure(png, caption_label, caption_body, width_cm=16.0):
    path = FIGS_MAIN / png if (FIGS_MAIN / png).exists() else FIGS_SUPP / png
    if not path.exists():
        add_md(f'[FIGURE NOT FOUND: {png}]', italic=True); return
    img = Image.open(path); w, h = img.size
    if width_cm * h / w > 21.0:
        width_cm = 21.0 * w / h
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(caption_label, caption_body)


def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _cell(cell, text, bold=False, size=8.5, bg=None, center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.font.name = 'Arial'; run.font.size = Pt(size); run.bold = bold
    if bg:
        _shade(cell, bg)


def _table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'bottom']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8'); b.set(qn('w:color'), '000000')
        borders.append(b)
    for side in ['left', 'right', 'insideV']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'none'); borders.append(b)
    ih = OxmlElement('w:insideH'); ih.set(qn('w:val'), 'single'); ih.set(qn('w:sz'), '2'); ih.set(qn('w:color'), 'CCCCCC')
    borders.append(ih)
    tblPr.append(borders)


def add_table(headers, rows, label, caption):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(tbl)
    for j, h in enumerate(headers):
        _cell(tbl.rows[0].cells[j], h, bold=True, bg='1A237E', size=8.5)
        # white header text
        tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            _cell(tbl.rows[i + 1].cells[j], str(val), size=8.5,
                  bg=('EEF2FF' if i % 2 else 'FFFFFF'))
    add_caption(label, caption)
    doc.add_paragraph()


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
set_run_font(p.add_run(
    'PAD4-DB: a curated structure–activity resource reveals hub-organized activity '
    'cliffs and scaffold-dependent SAR ruggedness in PAD4 inhibitors'), size=16, bold=True)

for txt, sz in [
    ('[Author 1]¹, [Author 2]², [Corresponding Author]¹*', 11),
    ('¹ [Affiliation 1]   ² [Affiliation 2]', 10),
    ('* Correspondence: [email]', 10),
]:
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(q.add_run(txt), size=sz, italic=True)

doc.add_paragraph()
p = doc.add_paragraph()
set_run_font(p.add_run('Keywords: '), size=10, bold=True)
set_run_font(p.add_run('PAD4; PADI4; protein arginine deiminase; activity cliffs; matched molecular pairs; '
                       'structure–activity relationship; cheminformatics; database'), size=10)
p = doc.add_paragraph()
set_run_font(p.add_run('Abbreviations: '), size=10, bold=True)
set_run_font(p.add_run('PAD4, protein arginine deiminase 4; SAR, structure–activity relationship; '
                       'HTS, high-throughput screening; MMP, matched molecular pair; SALI, structure–activity '
                       'landscape index; ECFP, extended-connectivity fingerprint; SAS, structure–activity '
                       'similarity; pIC50, −log10(IC50 in M).'), size=10)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Abstract', 1)
add_md(
    'Protein arginine deiminase 4 (PAD4/PADI4) is a calcium-dependent enzyme implicated in rheumatoid '
    'arthritis, NETosis-driven pathologies, and oncology, and is an actively pursued drug target. Public '
    'bioactivity data for PAD4 inhibitors are fragmented across heterogeneous assay formats and repositories, '
    'with inconsistent units, redundant deposits, and uncurated structures, which impedes quantitative '
    'structure–activity relationship (SAR) analysis and machine-learning model development. Here we present '
    '**PAD4-DB**, a standardized, deduplicated knowledge base assembled from PubChem bioassay campaigns, '
    'patent-deposited screening data, ChEMBL (assay CHEMBL6111) and BindingDB (UniProt Q9UM07) through a fully '
    'scripted, fail-loud pipeline. The resource comprises **3,093 structurally resolved inhibitors** with '
    'consensus pIC50 values (range 2.00–8.52; median 6.84) and a parallel high-throughput screening (HTS) layer '
    'of 327,336 screened compounds. Murcko scaffold analysis reveals strong chemotype concentration '
    '(Gini = 0.532; 71.9% of compounds in series of ≥2 members) and pronounced scaffold-dependent SAR '
    'ruggedness. Of 358,416 structurally related compound pairs, only 94 (0.026%) constitute severe activity '
    'cliffs (Tanimoto ≥ 0.8, |ΔpIC50| ≥ 2.0)—13-fold fewer than expected under permuted potencies '
    '(P < 0.001)—of which **80 (85.1%)** are corroborated by matched '
    'molecular pair (MMP) analysis. The severe-cliff network is organized around just **four hub compounds in '
    'two structurally and topologically distinct classes**, which together account for 50 of 94 severe cliff pairs '
    '(53.2%; 3.9-fold above a label-permutation null, P < 0.001) and whose dominance is invariant to fingerprint '
    'resolution (53.2% ECFP4, 53.3% ECFP6). '
    'PAD4-DB provides a transparent, reproducible foundation for PAD4 medicinal chemistry and an annotated '
    'benchmark of interpretable failure modes for similarity-based predictive models. All data, code, and '
    'figures are openly available.')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Introduction', 1)
add_md('Protein arginine deiminase 4 (PAD4, gene PADI4; UniProt Q9UM07) catalyzes the calcium-dependent '
       'conversion of peptidyl-arginine to peptidyl-citrulline, a post-translational modification that regulates '
       'chromatin decondensation, transcription, and the formation of neutrophil extracellular traps (NETs) [ref]. '
       'Dysregulated citrullination is mechanistically linked to rheumatoid arthritis, where citrullinated '
       'proteins are targets of disease-defining autoantibodies, and to NET-driven inflammation, thrombosis, and '
       'tumor progression [ref]. Consequently, PAD4 has attracted sustained drug-discovery interest, spanning '
       'covalent haloacetamidine warheads (e.g., Cl-amidine, F-amidine, BB-Cl-amidine) [ref], reversible '
       'benzimidazole series (e.g., GSK484, GSK199) [ref], and clinical-stage candidates [ref].')
add_md('Despite this activity, the public bioactivity record for PAD4 is fragmented. Measurements are distributed '
       'across primary repositories (PubChem, ChEMBL, BindingDB) that re-curate overlapping screening campaigns; '
       'they span incompatible assay formats (colorimetric BAEE hydrolysis, fluorescence-based RFMS, fluorescence '
       'polarization binding, and irreversible-inhibition kinetics); and they are reported in heterogeneous units '
       '(IC50, Ki, Kd, percent inhibition, kinact/KI). Raw deposits frequently contain salt forms, vendor-specific '
       'extended SMILES annotations, and duplicate or structure-less entries. These issues collectively obstruct '
       'three goals central to modern medicinal chemistry: (i) construction of a unit-consistent potency variable '
       'suitable for quantitative SAR; (ii) honest assessment of cross-source agreement, which is confounded by '
       're-curation rather than independent replication; and (iii) systematic mapping of the activity-cliff '
       'landscape, increasingly recognized as the principal obstacle to reliable machine-learning potency '
       'prediction [ref].')
add_md('Activity cliffs—pairs of structurally similar molecules with unexpectedly large potency differences—'
       'define the regions where the similarity principle breaks down and where similarity-based models fail most '
       'severely [ref]. Their organization (whether diffuse or concentrated around specific compounds or '
       'scaffolds) is of direct practical importance, yet has not been characterized for PAD4.')
add_md('Here we address these gaps with **PAD4-DB**, a curated, reproducible resource and an accompanying '
       'SAR-landscape analysis. We (1) describe a transparent pipeline that standardizes and deduplicates '
       'heterogeneous PAD4 bioactivity into 3,093 structurally resolved inhibitors with consensus pIC50 values; '
       '(2) quantify cross-source overlap and introduce a per-compound source-independence score that separates '
       'genuine replication from pipeline redundancy; (3) validate the resource against curated reference '
       'inhibitors; and (4) map the scaffold and activity-cliff landscape, revealing scaffold-dependent SAR '
       'ruggedness and a small set of hub compounds that dominate the cliff structure. We further demonstrate '
       'that the central findings are robust to fingerprint choice and are corroborated by an orthogonal '
       'substructure-based (MMP) method.')

# ══════════════════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Results', 1)

add_heading('A standardized, deduplicated PAD4 inhibitor resource', 2)
add_md('We assembled PAD4-DB from four source families: 57 confirmatory and 11 literature-derived PubChem '
       'bioassays (Layer A/C), 26 secondary PubChem assays (Layer D/E), three HTS campaigns (AIDs 463073, 485272, '
       '488796), ChEMBL assay CHEMBL6111, and the BindingDB record for UniProt Q9UM07 (Methods). A fully scripted, '
       'fail-loud pipeline ingested 341,282 measurement rows, standardized structures with RDKit, normalized '
       'activities to a unit-consistent scale, aggregated replicates, and deduplicated by InChIKey within source '
       'and endpoint type.')
add_md('Structure standardization succeeded for 341,276 of 341,282 rows (100.0%), with zero parse or '
       'sanitization failures after stripping BindingDB Daylight extended-SMILES annotations and salts (Methods). '
       'Activity normalization yielded 338,021 records with valid endpoints (99.0%); percent-inhibition HTS rows '
       '(the dominant endpoint, n = 330,136) were intercepted before unit conversion to prevent spurious nanomolar '
       'coercion. After replicate aggregation and deduplication, the resource resolves into two clearly separated '
       'layers: a **dose–response potency space of 3,093 unique inhibitors** carrying consensus pIC50 values, and '
       'a **HTS space of 327,336 unique screened compounds** carrying single-concentration percent-inhibition '
       'readouts (Figure 1). Two ChEMBL entries with qualifying IC50 values but no deposited structure were '
       'correctly excluded, yielding 3,095 qualifying measurements but 3,093 structure-resolved compounds.')
add_md('The 3,093 inhibitors span a broad potency range (pIC50 2.00–8.52; mean 6.55, median 6.84, SD 0.99) with '
       'a bimodal distribution: a main mode near pIC50 7.0 and a lower-potency shoulder near 5.0–6.0 driven by '
       'patent-exclusive screening hits (Figure 3a).')
add_figure('fig01_headline.png', 'Figure 1. Chemical-space landscape of the PAD4-DB potency set.',
           'Two-dimensional t-SNE embedding of ECFP4 fingerprints for all 3,093 inhibitors (grey points; soft '
           'density underlay). Severe activity-cliff compounds (n = 99) are colored by consensus pIC50 (viridis); '
           'the four cliff-hub compounds are marked (Class A, navy stars; Class B, red diamonds). t-SNE was used '
           'for visualization only, not for quantitative analysis. Fingerprints: Morgan radius 2, 2,048 bits '
           '(RDKit 2025.09.5).')

add_heading('Cross-source overlap reflects re-curation, not independent replication', 2)
add_md('A naïve reading of source agreement would overstate data quality. 2,755 of 3,093 compounds (89.1%) are '
       'multi-source, and pairwise potency concordance is 99.7% (3,084/3,093 concordant; 0 discordant). However, '
       'this structure arises because the primary repositories re-curate overlapping assay campaigns rather than '
       'reporting independent experiments (Figure 2a,b). The two dominant redundancy patterns—'
       'BindingDB+ChEMBL+PubChem (n = 1,366) and BindingDB+PubChem (n = 1,199)—together cover 2,565 of 3,093 '
       'compounds (82.9%).')
add_md('To distinguish genuine replication from pipeline redundancy, we assigned each compound a source-'
       'independence score penalizing known re-curation links (Methods). Under this score, only **528 compounds '
       '(17.1%) are non-redundant** (score ≥ 0.6), while 2,565 (82.9%) are flagged as pipeline-redundant '
       '(Figure 2c). We therefore report the source-independence score, not the raw multi-source fraction, as the '
       'basis for confidence weighting.')
add_figure('fig02_source_overlap.png', 'Figure 2. Source composition and independence.',
           '(a) UpSet plot of source-combination membership across PubChem, ChEMBL and BindingDB. (b) Per-source '
           'coverage. (c) Source-independence partition: 528 non-redundant (17.1%) versus 2,565 pipeline-redundant '
           'compounds (82.9%). Independence score derived from source-combination re-curation links (Methods); '
           'threshold ≥ 0.6 defines non-redundant.')
add_md('**Interpretation.** Cross-source overlap reflects partial re-curation of shared primary assay campaigns '
       'rather than fully independent experimental replication. Concordance (99.7%) reflects this structure, not '
       'measurement agreement.', italic=True)

add_heading('Potency is structured by assay mechanism', 2)
add_md('Consensus potency is statistically indistinguishable across the major repositories (PubChem mean pIC50 '
       '6.62; BindingDB 6.59; ChEMBL 6.50; medians 6.85/6.85/6.90; Kruskal–Wallis H = 0.25, P = 0.88; Table 1) '
       'but is strongly differentiated by assay mechanism (Figure 3c). '
       'Fluorescence-polarization binding assays report the highest mean potency (pIC50 7.18, n = 115), enzymatic '
       'RFMS-confirmed measurements are intermediate (6.72, n = 878), colorimetric BAEE assays form the bulk '
       '(6.47, n = 2,079), and covalent-mechanism measurements are lowest and most dispersed (4.21, n = 21), '
       'consistent with single-timepoint IC50 values underestimating the potency of time-dependent irreversible '
       'inhibitors [ref]. Patent-exclusive compounds (n = 233) differ from published compounds principally in '
       'distribution shape rather than central tendency: their median pIC50 is nearly identical (Δmedian = 0.04) '
       'but a heavier low-potency tail lowers their mean (6.13 versus 6.53), broadening the low-potency shoulder '
       '(Figure 3a,b; Mann–Whitney U P = 1.3 × 10⁻⁵, rank-biserial r = 0.17, a small effect). '
       'Across the chemistry, 107 compounds (3.5%) carry '
       'electrophilic warheads (chloroacetamidine n = 66, fluoroacetamidine n = 17, haloacetyl n = 11, and minor '
       'enaminone, vinyl-sulfone, and α-bromoketone classes), flagged by SMARTS pattern matching (Methods). '
       'Consistent with the beyond-rule-of-five chemotypes that dominate PAD4 chemistry, potency correlates '
       'moderately with molecular size (pIC50 versus molecular weight, Spearman ρ = 0.54; versus heavy-atom count, '
       'ρ = 0.57; both P < 10⁻²²⁰); we report this size–potency trend as a deposited annotation, as it is a '
       'recognized confound for potency models and should be controlled for in predictive use of the resource.')
add_figure('fig03_potency.png', 'Figure 3. Potency landscape.',
           '(a) Consensus pIC50 distribution (histogram + kernel density; mean 6.55, median 6.84, SD 0.99; '
           'n = 3,093). (b) pIC50 by source and patent status (violin plots; medians marked). (c) Compound counts '
           'and mean pIC50 by assay-mechanism class. Mann–Whitney U test, two-sided, PubChem vs patent-only, '
           'P < 0.001; n shown per group.')

add_heading('Validation against curated reference inhibitors', 2)
add_md('To confirm that PAD4-DB faithfully recovers known pharmacology, we audited 14 curated reference '
       'inhibitors (Table 2). Seven are present with consensus pIC50 values concordant with the literature '
       '(mean |ΔpIC50| < 0.15 log units), including GSK484 (7.049), BMS-P5 (7.009), TDFA (5.638), '
       'Streptonigrin (5.602), Cl-amidine (5.219), F-amidine (4.571), and JBI-589 (6.000). Three reference '
       'compounds are present in the raw data but correctly excluded from the dose–response set because they lack '
       'a primary IC50 endpoint (o-F-amidine, kinact/KI only; Amodiaquine, HTS percent-inhibition only; '
       'BB-Cl-amidine, covalent kinetics only). Three are absent from all source databases (GSK199, Pyroxamide, '
       'PAD-PF1), and one (AFM-30a) is correctly absent as a PAD2-selective compound. A salt-form discrepancy for '
       'GSK484 was resolved correctly by the standardization pipeline, which used the free base rather than the '
       'deposited hydrochloride form.')
add_md('**Scope.** PAD4-DB comprises compounds from PubChem bioassay campaigns, patent-deposited screening data, '
       'ChEMBL (CHEMBL6111), and BindingDB (Q9UM07). Seven of thirteen curated PAD4 reference inhibitors are '
       'present with concordant pIC50 values (mean |ΔpIC50| < 0.15 log units). Three additional compounds are '
       'present but lack primary IC50 measurements (covalent kinetics or HTS-only data). Three compounds are '
       'absent from all source databases, reflecting gaps in public bioactivity curation rather than pipeline '
       'exclusions.', italic=True)

add_heading('Scaffold architecture and SAR ruggedness', 2)
add_md('The 3,093 inhibitors map to **1,244 unique Bemis–Murcko scaffolds**, comprising 375 multi-member series '
       'and 869 singletons. Chemotype representation is highly concentrated (Gini coefficient 0.532): the single '
       'largest series—an azaindole–benzimidazole framework—contains 174 compounds, and 71.9% of all compounds '
       'belong to series of ≥2 members (Figure 4a,b). Patent-deposited screening contributes genuinely novel '
       'chemical space: 103 scaffolds are exclusive to patent-derived compounds, at series density identical to '
       'the published set (mean 2.5 members) but at modestly lower potency (mean pIC50 6.13 versus 6.53).')
add_md('Crucially, intra-scaffold potency spread is large and heterogeneous, providing direct evidence of '
       '**scaffold-dependent SAR ruggedness** (Figure 4c): the median within-series pIC50 standard deviation is '
       'σ = 0.27 log units, but individual series reach σ > 1.0, and the dominant 174-member series '
       'itself spans σ = 0.45 with a pIC50 range of 2.66 log units. Ruggedness scales with optimization effort: '
       'across the 375 multi-member series, series size correlates positively with intra-scaffold σ '
       '(Spearman ρ = 0.36, P = 5 × 10⁻¹³), so the most heavily elaborated chemotypes are also the most rugged. '
       'This ruggedness is the structural '
       'substrate of the activity cliffs characterized below; cliff density is concentrated in a minority of '
       'series (Figure 4d).')
add_figure('fig04_scaffold.png', 'Figure 4. Scaffold landscape and SAR ruggedness.',
           '(a) Top-15 Murcko scaffold series by size, colored by mean pIC50; patent-exclusive scaffolds outlined. '
           '(b) Lorenz curve of scaffold-size distribution (Gini = 0.532). (c) SAR ruggedness: series size (log) '
           'versus intra-scaffold pIC50 spread (σ); median σ = 0.27 marked. (d) Severe-cliff density '
           '(cliff pairs / possible pairs) for series of ≥4 members. Scaffolds computed as generic '
           'Bemis–Murcko frameworks (RDKit 2025.09.5).')

add_heading('The activity-cliff landscape is sparse and quadrant-structured', 2)
add_md('We computed all pairwise ECFP4 Tanimoto similarities and potency differences for the 358,416 structurally '
       'related pairs (Tanimoto ≥ 0.6). The structure–activity similarity (SAS) map (Supplementary '
       'Figure S10) is dominated by smooth SAR: 96.06% of related pairs are non-descript (low similarity, low '
       'ΔpIC50), 3.34% are continuous/smooth SAR (high similarity, low ΔpIC50), 0.57% are discontinuous '
       '(low similarity, high ΔpIC50), and only **0.026% (94 pairs) are severe activity cliffs** '
       '(Tanimoto ≥ 0.8 and |ΔpIC50| ≥ 2.0). Even among near-identical pairs (Tanimoto 0.9–1.0), '
       'only 0.61% exceed the cliff threshold. This scarcity is not a trivial consequence of the potency range: '
       'permuting pIC50 labels across the 2,620 compounds of the high-similarity (Tanimoto ≥ 0.8) subgraph yields '
       '1,225 ± 100 severe cliffs (10,000 permutations), so the observed 94 represent a **13-fold depletion '
       'relative to chance (P < 0.001)**—a quantitative confirmation of pronounced “diagonal absence” and of a '
       'globally smooth landscape [ref].')
add_md('The 2.0-log severe-cliff threshold is also well above the measurement noise floor: the maximum '
       'cross-source pIC50 spread across all multi-source compounds is 0.74 log units (median 0), so severe cliffs '
       'exceed the largest observed inter-repository disagreement by 2.7-fold and cannot be attributed to '
       'measurement error. Applying graded thresholds yields 94 severe (|ΔpIC50| ≥ 2.0), 193 moderate '
       '(1.5 ≤ |ΔpIC50| < 2.0), and 580 broad (1.0 ≤ |ΔpIC50| < 1.5) cliff pairs, involving 99, 209, and '
       '539 compounds respectively (Table 3). The maximum severe-cliff potency difference is 3.045 log units '
       '(≈1,100-fold), and the mean severe |ΔpIC50| is 2.31.')
add_figure('fig_s03_sas_map.png', 'Supplementary Figure S3. Structure–activity similarity (SAS) map.',
           '(a) Pairwise ECFP4 Tanimoto similarity versus |ΔpIC50| for all 358,416 related pairs (hexbin, '
           'log density); 94 severe cliffs highlighted in the upper-right activity-cliff quadrant. (b) |ΔpIC50| '
           'distribution stratified by similarity bin with per-bin cliff rate, demonstrating diagonal absence '
           '(only 0.61% of near-identical pairs exceed the cliff threshold). Cliff thresholds (Tanimoto = 0.8, '
           '|ΔpIC50| = 2.0) follow Senger (2009) and Stumpfe & Bajorath (2012) [ref].')

add_heading('Severe activity cliffs are organized around four hub compounds', 2)
add_md('The severe-cliff network (99 nodes, 94 edges) is not diffuse but is dominated by a small number of '
       'high-degree hubs (Figure 5). Four compounds, falling into two structurally and topologically distinct '
       'classes, together account for **50 of 94 severe cliff pairs (53.2%)** (Table 4). This concentration is far '
       'above chance: under a label-permutation null that shuffles pIC50 values across the high-similarity subgraph '
       'while holding its structure fixed, the top-four compounds account for only 13.6 ± 2.5% of severe cliffs, '
       'so the observed 53.2% represents a **3.9-fold enrichment (P < 0.001, 10,000 permutations)**. The hub '
       'structure is therefore potency-driven, not an artifact of which compounds happen to have many close '
       'structural neighbors.')
add_md('**Class A — series-embedded mid-potency floors.** Two compounds (SMADULGDNOCLOP-GISFHXKWSA-N, pIC50 5.39, '
       '15 cliff pairs; RAVBZQAQTVGKIV-XBPDSQQVSA-N, pIC50 5.34, 12 cliff pairs) are mid-potency members of the '
       'dominant 174-member azaindole–benzimidazole series. Their position as within-series potency floors '
       'generates 27 severe cliff pairs against higher-potency analogs of the same chemotype.')
add_md('**Class B — scaffold-singleton structural attractors.** Two compounds (UDCDEKJNAMHBFH-HSZRJFAPSA-N and '
       'DVCKJOQIVOGXEI-XMMPIXPASA-N; both pIC50 4.30) carry distinct cyclobutyl- versus cyclopentyl-sulfonamide '
       'rings. Because the differing ring is part of the ring system retained by the Bemis–Murcko framework, each '
       'compound is assigned a unique (singleton) Murcko scaffold even though their whole-molecule ECFP4 similarity '
       'is high (mutual Tanimoto 0.975, a single-methylene difference); this combination of scaffold uniqueness and '
       'high fingerprint similarity is precisely what makes them broad structural attractors, generating 23 severe '
       'cliff pairs spanning multiple chemotypes. Their shared free primary amine is a plausible source of assay '
       'interference or early-stage, unoptimized potency.')
add_md('Cliff-hub identity is potency-defined: hubs are markedly less potent than the other 95 severe-cliff '
       'compounds (Δmedian = 2.4 log units; Mann–Whitney U P = 0.007; rank-biserial r = 0.80, a large effect). '
       'We detected no accompanying differences on any physicochemical descriptor examined (molecular weight, '
       'cLogP, TPSA, hydrogen-bond donors/acceptors, rotatable bonds, aromatic rings, fraction Csp3; all P > 0.26; '
       'Supplementary Table S-hub); however, with only four hubs this comparison is underpowered, so we interpret '
       'it as the absence of any large physicochemical distinction rather than proof of equivalence. The hubs are '
       'best understood as otherwise unremarkable molecules whose cliff behavior derives from their landscape '
       'position, not from anomalous physicochemistry.')
add_figure('fig05_cliff_network.png', 'Figure 5. Severe activity-cliff network.',
           '(a) Network of 99 compounds (nodes) connected by 94 severe cliff edges; node size ∝ degree, '
           'node color = consensus pIC50, edge color = |ΔpIC50|. Class A hubs (navy stars) and Class B hubs '
           '(red diamonds) are highlighted. (b) Top-12 compounds by severe-cliff degree (non-hub bars labelled by '
           'InChIKey skeleton). The four hubs account for 50/94 severe cliff pairs (3.9-fold above a '
           'label-permutation null, P < 0.001).')

add_heading('Cliffs are corroborated by MMP analysis and are fingerprint-robust', 2)
add_md('To ensure that severe cliffs are not artifacts of fingerprint choice, we applied two orthogonal '
       'validations. First, matched molecular pair (MMP) analysis confirmed **80 of 94 severe cliff pairs '
       '(85.1%)** through an explicit shared chemical core (Figure 6); these decompose into 45 single-atom changes '
       '(56%), 27 small-substituent changes (34%), and 8 medium-substituent changes (10%), demonstrating that '
       'large potency discontinuities frequently arise from minimal structural perturbation. Second, sensitivity '
       'analysis with ECFP6 (radius 3) showed that 64 of 94 pairs have ECFP4 Tanimoto 0.80–0.85 and fall below '
       '0.80 at radius 3—expected behavior for large fused-ring systems—yet 80% of these borderline pairs (51/64) '
       'remain MMP-confirmed, and **hub dominance is fingerprint-invariant (53.2% under ECFP4, 53.3% under ECFP6)**. '
       'Thirteen pairs (13.8%) are severe by ECFP4 alone without MMP or ECFP6 corroboration; these are explicitly '
       'flagged (ecfp4_only_cliff = True) in the deposited dataset. As a further control, severe-cliff compounds '
       'are not enriched in any single assay-mechanism class relative to the full dataset (Fisher’s exact test, '
       'all P > 0.05; Supplementary Figure S2), arguing against an assay-format origin for the cliffs.')
add_figure('fig06b_cliff_pairs.png', 'Figure 6. Matched molecular pair analysis of severe cliffs.',
           '(a) MMP-confirmed severe cliff pairs by change type (n = 80: 45 single-atom, 27 small-substituent, '
           '8 medium-substituent). (b) Four representative cliff pairs with the higher-potency (gain, blue) and '
           'lower-potency (loss, orange) difference atoms highlighted on the more- and less-potent analog '
           'respectively. Colourblind-safe blue/orange highlighting; MMP cores '
           'derived by seeded maximum-common-substructure search (rdFMCS); pairs with disconnected difference '
           'fragments excluded.')

add_heading('Medicinal-chemistry interpretation of the cliff hubs', 2)
add_md('The four cliff-hub compounds are not merely a statistical feature; they map onto two concrete structural '
       'vulnerabilities in the PAD4 inhibitor landscape. Class A hubs (SMADULGDNOCLOP-GISFHXKWSA-N and '
       'RAVBZQAQTVGKIV-XBPDSQQVSA-N) act as within-series potency floors inside the heavily optimized 174-member '
       'azaindole–benzimidazole chemotype. This is a familiar but underappreciated hazard in lead optimization: '
       'as a series is elaborated, individual substituent changes can open deep potency valleys—low-activity '
       'analogs that remain nearly identical, by both fingerprint and matched molecular pair, to high-potency '
       'leads. Single-atom MMP transformations account for 56% of severe cliffs, underscoring that such valleys '
       'are reachable in one synthetic step.')
add_md('Class B hubs (UDCDEKJNAMHBFH-HSZRJFAPSA-N and DVCKJOQIVOGXEI-XMMPIXPASA-N) present the opposite problem. '
       'As scaffold singletons sharing a free primary amine and a uniformly low potency (pIC50 = 4.30), they are '
       'plausibly reporting a systematic assay liability—for example, ionic interference or non-specific behavior '
       'in the enzymatic readout—rather than genuine, structure-driven PAD4 inhibition. Because their ECFP4 '
       'representation is broadly similar to several disparate chemotypes, they behave as structural attractors '
       'that would bias any similarity-based predictor toward overestimating the potency of their neighbors.')
add_md('Together these archetypes furnish a compact, mechanistically grounded benchmark. A model that can place '
       'Class A floor compounds correctly relative to their potent within-series neighbors, and that resists the '
       'Class B attractors, is demonstrating that it has learned target-relevant SAR rather than fingerprint '
       'proximity. We therefore deposit per-compound hub_class labels (A/B/none) to enable cliff-aware model '
       'evaluation; whether these compounds constitute actual failure modes for a given architecture is itself an '
       'empirical question the resource is designed to support.')

add_heading('Parallel HTS layer', 2)
add_md('The 327,336-compound HTS layer provides an orthogonal, single-concentration screening context. 1,453 '
       'dose–response compounds share InChIKeys with HTS-screened compounds, indicating parallel measurement in '
       'independent assay pipelines. Only 6 of these were confirmed HTS actives (≥50% inhibition at screening '
       'concentration); the remaining 1,447 carry literature IC50 values yet show low HTS inhibition (median '
       '4.3%), consistent with potent compounds (median pIC50 6.93) tested below their IC50 at standard HTS '
       'concentrations.')

# ══════════════════════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Discussion', 1)
add_md('PAD4-DB converts a fragmented public record into a transparent, reproducible SAR resource and, in doing '
       'so, exposes the structure of the PAD4 inhibitor landscape. Beyond cataloguing what the landscape contains, '
       'the analysis lets us ask why it has the shape it does.')
add_md('A first, methodological point concerns data provenance: **apparent cross-source agreement is largely an '
       'artifact of re-curation**. The 99.7% concordance '
       'and 89.1% multi-source fraction would, if taken at face value, suggest near-perfect independent '
       'replication; the source-independence score reveals that only 17.1% of compounds are genuinely '
       'non-redundant. This distinction matters for any downstream confidence weighting or uncertainty estimate, '
       'and we recommend reporting independence-aware statistics rather than raw multi-source counts for '
       'aggregated bioactivity resources generally.')
add_md('**Why does PAD4 produce cliff hubs?** The landscape is globally smooth (severe cliffs are 13-fold rarer '
       'than chance) yet locally concentrated, with four compounds accounting for 53.2% of severe cliffs (3.9-fold '
       'above null; both P < 0.001). We propose that this organization follows directly from the geometry of the '
       'PAD4 active site. PAD4 buries its substrate arginine in a narrow, highly electronegative, calcium-gated '
       'pocket terminating at the catalytic Cys645–His471 dyad [ref]; potency in the dominant reversible chemotypes '
       'therefore depends on a small number of make-or-break contacts to this constrained region. Class A hubs are '
       'the structural manifestation of this sensitivity: as mid-potency members of the heavily elaborated '
       'azaindole–benzimidazole series, they sit one substituent removed—frequently a single atom (single-atom MMP '
       'transformations underlie 56% of severe cliffs)—from analogs that recover the critical contact and gain 2–3 '
       'log units. They are not anomalous molecules but the lower rim of a steep local SAR around a conserved '
       'pharmacophore. Class B hubs arise from a different mechanism: a shared free primary amine confers uniformly '
       'low, assay-format-sensitive potency (consistent with electrostatic or buffer-dependent behavior against the '
       'anionic pocket rather than defined active-site engagement), while their broad ECFP4 resemblance to many '
       'amine-bearing chemotypes makes them promiscuous structural attractors. The two hub classes thus encode two '
       'genuinely different phenomena—a steep but real pharmacophore boundary (Class A) and a likely assay-driven '
       'artifact (Class B)—and together they constitute a compact, mechanistically interpretable benchmark of '
       'failure modes for similarity-based models.')
add_md('**Why do only certain scaffolds generate cliffs?** Cliffs are not a generic property of large series but '
       'require the coincidence of two conditions: dense local sampling (so that near-neighbors exist) and an '
       'intrinsically steep SAR vector (so that small changes matter). Only 11 of 155 scaffold series with ≥4 '
       'members harbor any within-scaffold severe cliff (Supplementary Fig. S1), and ruggedness rises with '
       'optimization effort (series size versus intra-scaffold σ, Spearman ρ = 0.36, P = 5 × 10⁻¹³). Yet sampling '
       'density alone is insufficient: the second-largest series (n = 102) is entirely smooth (σ = 0.37, zero '
       'cliffs) despite being sampled as densely as the rugged 174-member series that contains the Class A hubs. '
       'The decisive factor is therefore where a series varies—chemotypes that diversify on solvent-exposed or '
       'tolerant vectors remain smooth, whereas those that vary at a position contacting the catalytic or '
       'calcium-binding machinery become rugged. SAR ruggedness in PAD4 is consequently scaffold- and '
       'vector-specific, and predictive models should expect heteroscedastic difficulty: high confidence on smooth '
       'families, but poorly calibrated extrapolation within the few rugged ones.')
add_md('**Implications for future PAD4 optimization campaigns.** These observations translate into concrete '
       'guidance. First, within the azaindole–benzimidazole chemotype, single-atom modifications at the '
       'cliff-forming vector carry disproportionate risk; campaigns should map this vector early and protect the '
       'substituent that distinguishes potent analogs from the Class A floor rather than treating it as freely '
       'optimizable. Second, free-amine singletons resembling the Class B hubs should be treated as assay-liability '
       'risks rather than SAR signal—capping the amine or confirming activity in an orthogonal (e.g., '
       'non-fluorogenic) format is warranted before such compounds inform a series. Third, the per-scaffold '
       'cliff-density ranking (Supplementary Fig. S1, Supplementary Table S-scaf) provides a triage map: rugged '
       'scaffolds demand denser analoging and mechanism-aware quality control, whereas smooth scaffolds are '
       'comparatively safe substrates for property optimization. Finally, the moderate size–potency correlation '
       '(ρ = 0.54) cautions that apparent potency gains driven by increasing molecular weight should be monitored '
       'by ligand efficiency, particularly given the beyond-rule-of-five regime these inhibitors already occupy.')
add_md('Several limitations bound interpretation. PAD4-DB is **assay-registry-derived, not a comprehensive '
       'medicinal-chemistry reference space**; absences (e.g., GSK199, Pyroxamide, PAD-PF1) reflect gaps in public '
       'curation rather than pipeline errors. Potency values for time-dependent covalent inhibitors are reported '
       'as deposited single-timepoint IC50 and underestimate true inhibitory efficiency; these compounds are '
       'flagged but should be modeled with mechanism-aware care. Finally, consensus pIC50 values aggregate '
       'measurements across assay formats with differing calcium concentrations and substrates, which can '
       'introduce systematic offsets (e.g., the 0.9-log JBI-589 discrepancy traced to differing assay conditions, '
       'not curation error).')
add_md('These caveats notwithstanding, PAD4-DB offers what the field has lacked: a single, standardized, '
       'deduplicated, and openly documented PAD4 inhibitor set with explicit provenance, confidence annotation, '
       'and a characterized SAR landscape.')

# ══════════════════════════════════════════════════════════════════════════════
# METHODS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Methods', 1)

add_heading('Data sources and target identity', 3)
add_md('PAD4-DB was assembled for human PAD4 (gene PADI4; UniProt Q9UM07). Sources comprised: 57 confirmatory and '
       '11 literature-derived PubChem bioassays; 26 secondary PubChem bioassays; three HTS campaigns (AIDs 463073, '
       '485272, 488796); the ChEMBL bioactivity export for assay CHEMBL6111; and the BindingDB record for Q9UM07. '
       'Target identity was audited per assay: 93 assays were explicitly PAD4, two were PAD-family, and two were '
       'ambiguous; all non-explicit assays contributed zero dose–response records. Species filtering retained '
       'Homo sapiens (ChEMBL: 4,858 human, 67 unknown, 0 non-human; BindingDB: 3,087 human, 0 non-human). PubChem '
       'bioassay data were downloaded 2026-06-10 to 2026-06-14; ChEMBL (CHEMBL6111) on 2026-06-14; BindingDB '
       '(Q9UM07) on 2026-06-10.')

add_heading('Structure standardization', 3)
add_md('Structures were processed with RDKit 2025.09.5 (Python 3.10.19, pandas 2.3.3, numpy 2.2.5). BindingDB '
       'exports include Daylight extended-SMILES annotations that RDKit cannot parse; these suffixes were stripped '
       'prior to MolFromSmiles. Salts were removed and parent structures retained; InChIKeys were computed for '
       'deduplication. Of 341,282 ingested rows, 341,276 standardized successfully (100.0%; 6 NO_SMILES, 0 parse '
       'failures, 0 sanitization failures), yielding 328,976 unique InChIKeys across all layers.')

add_heading('Activity normalization, aggregation, and deduplication', 3)
add_md('Endpoints were normalized and, where an IC50-type concentration was available, converted to '
       'pIC50 = −log10(IC50 [M]). Percent-inhibition HTS rows (n = 330,136) were intercepted before nanomolar '
       'conversion (raising endpoint validity from 89.8% to 99.0%). ChEMBL pChEMBL values were cross-checked '
       '(0 mismatches). Measurements were grouped by InChIKey × source × assay ID × endpoint type '
       'and consolidated by log-mean aggregation (450 multi-replicate groups; maximum within-group discrepancy '
       '= 0.0), then deduplicated to InChIKey × source × endpoint level. This separates the data into a '
       'dose–response potency space (3,093 structure-resolved compounds) and an HTS space (327,336 compounds). '
       'Two ChEMBL compounds with qualifying IC50 but no deposited structure were excluded.')

add_heading('Source-independence scoring', 3)
add_md('Each compound’s source-combination membership was mapped to an independence score that penalizes '
       'known re-curation links among PubChem, ChEMBL, and BindingDB (single-source = 1.0; independent '
       'multi-source combinations score higher than re-curated ones). Compounds with score ≥ 0.6 were '
       'classified as non-redundant. The score distribution was 0.3 (n = 1,366), 0.5 (n = 1,199), 0.6 (n = 167), '
       '0.7 (n = 23), and 1.0 (n = 338).')

add_heading('Scaffold, fingerprint, and activity-cliff analysis', 3)
add_md('Generic Bemis–Murcko scaffolds were computed with RDKit. Structural similarity used ECFP4 fingerprints '
       '(Morgan algorithm, radius 2, 2,048 bits), following Senger (2009) and Stumpfe & Bajorath (2012) [ref]. '
       'All pairwise similarities were computed; the 358,416 pairs with Tanimoto ≥ 0.6 were retained. '
       'Activity cliffs were defined as pairs with Tanimoto ≥ 0.8 and |ΔpIC50| ≥ 2.0 (severe), with '
       'graded thresholds at |ΔpIC50| ≥ 1.5 (moderate) and ≥ 1.0 (broad). The Structure–Activity '
       'Landscape Index (SALI = |ΔpIC50| / (1 − Tanimoto)) was computed for all pairs. The cliff network '
       'used severe cliff pairs as edges; node degree defined hub compounds. Sensitivity analysis recomputed '
       'similarities with ECFP6 (radius 3): 30 pairs were robust (≥ 0.8), 64 non-robust (ECFP4 0.80–0.85, '
       'ECFP6 < 0.80) of which 51 (80%) remained MMP-confirmed; hub dominance was 53.2% (ECFP4) and 53.3% (ECFP6); '
       '13 ECFP4-only pairs were flagged.')

add_heading('Matched molecular pair analysis', 3)
add_md('For severe cliff pairs, a shared chemical core was identified by seeded maximum-common-substructure search '
       '(rdFMCS) with element and bond-order matching, and pairs were classified by changing-fragment size '
       '(single-atom, small-, medium-, large-substituent). Joining the canonical 94-severe-cliff set against the '
       'MMP table by canonical pair key, **80 of 94 severe pairs (85.1%)** yielded a valid connected MMP core '
       '(45 single-atom, 27 small-substituent, 8 medium-substituent). Difference atoms were colored by potency '
       'rank for visualization (gain = more-potent analog, loss = less-potent analog).')

add_heading('Covalent annotation and statistics', 3)
add_md('Electrophilic warheads were flagged by SMARTS matching (chloroacetamidine, fluoroacetamidine, haloacetyl, '
       'enaminone, vinyl-sulfone, α-bromoketone), identifying 107 covalent compounds (3.5%). Two-distribution '
       'comparisons used the two-sided Mann–Whitney U test, reported with the rank-biserial correlation as an '
       'effect size; multi-group comparison across sources used the Kruskal–Wallis test; categorical enrichment '
       'used Fisher’s exact test with the phi coefficient; monotonic associations used the Spearman rank '
       'correlation. Two label-permutation tests (10,000 iterations each, fixed seed) assessed the activity-cliff '
       'landscape on the Tanimoto ≥ 0.8 subgraph (2,620 compounds, 12,071 eligible pairs): consensus pIC50 values '
       'were permuted across subgraph compounds while the similarity structure was held fixed, and each iteration '
       'recorded the number of severe cliffs (testing cliff rarity) and the fraction incident to the four '
       'highest-degree compounds (testing hub concentration); empirical P values are reported as (k + 1)/(N + 1). '
       'The measurement noise floor was estimated from the per-compound cross-source pIC50 spread. All inferential '
       'statistics are reproduced by supp_statistical_tests.py and deposited in Supplementary Table S-stat. '
       'All analyses ran from a fixed project root under a pinned conda environment, with a '
       'pre-flight validation script asserting canonical counts (3,093 compounds; 94 severe cliffs; 4 hub '
       'compounds; 1,244 scaffolds; 707 MMP relationships among cliff compounds; 24 unique shared cores) before '
       'figure generation. Figures were rendered at 600 dpi following Nature Methods/Scientific Data standards.')

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Conclusion', 1)
add_md('PAD4-DB v2 is a rigorously curated PAD4 inhibitor resource that goes beyond data aggregation to reveal '
       'activity-cliff organization, hub-compound behavior, scaffold-dependent SAR ruggedness, and '
       'source-dependent chemical diversity. By standardizing 3,093 inhibitors with explicit provenance and '
       'confidence annotation, distinguishing genuine replication from pipeline redundancy, and mapping a sparse '
       'but hub-concentrated cliff landscape that is robust to fingerprint choice and corroborated by matched '
       'molecular pairs, the resource provides both a practical foundation for PAD4 medicinal chemistry and an '
       'interpretable benchmark of failure modes for similarity-based predictive models. All data, code, and '
       'figures are released to enable transparent reuse and extension.')

add_heading('Data and code availability', 3)
add_md('The complete PAD4-DB dataset (3,093-compound potency table with consensus pIC50, source-independence '
       'scores, scaffold and warhead annotations, and hub-class labels), the 327,336-compound HTS layer, the '
       'activity-cliff tables, and all analysis and figure-generation scripts are openly available at '
       '[repository DOI / URL — to be assigned]. Source bioactivity data derive from PubChem, ChEMBL '
       '(CHEMBL6111), and BindingDB (Q9UM07).')

# ══════════════════════════════════════════════════════════════════════════════
# TABLES
# ══════════════════════════════════════════════════════════════════════════════
page_break()
add_heading('Tables', 1)

add_table(
    ['Source', 'n', '% of 3,093', 'Mean pIC50', 'Median pIC50', 'SD'],
    [['PubChem (confirmatory)', '2,821', '91.2', '6.62', '6.85', '0.90'],
     ['BindingDB', '2,827', '91.4', '6.59', '6.85', '0.94'],
     ['ChEMBL', '1,566', '50.6', '6.50', '6.90', '1.10'],
     ['All three sources', '1,366', '44.2', '6.64', '6.94', '0.93']],
    'Table 1. Source coverage and consensus potency.',
    'Percentages sum to >100% because most compounds appear in multiple sources (89.1% multi-source). '
    'pIC50 = −log10(IC50 [M]); consensus across replicate measurements.')

add_table(
    ['Status', 'n', 'Compounds / notes'],
    [['Present, concordant', '7', 'Streptonigrin (5.602), Cl-amidine (5.219), F-amidine (4.571), '
      'GSK484 (7.049), TDFA (5.638), BMS-P5 (7.009), JBI-589 (6.000); mean |ΔpIC50| < 0.15 vs literature'],
     ['Present, not mapped', '3', 'o-F-amidine; Amodiaquine; BB-Cl-amidine (no primary IC50 endpoint)'],
     ['Absent by design', '3', 'GSK199; Pyroxamide; PAD-PF1 (not deposited in any source)'],
     ['Correctly excluded', '1', 'AFM-30a (PAD2-selective)']],
    'Table 2. Recovery of curated PAD4 reference inhibitors (n = 14).',
    'Values in parentheses are consensus pIC50. GSK484 recovered as free base after salt stripping.')

add_table(
    ['Tier', '|ΔpIC50| band', 'Pairs', 'Compounds', '% of dataset', 'Median |ΔpIC50|', 'Max |ΔpIC50|'],
    [['Severe', '≥ 2.0', '94', '99', '3.2', '2.23', '3.045'],
     ['Moderate', '1.5 – <2.0', '193', '209', '6.8', '1.70', '1.99'],
     ['Broad', '1.0 – <1.5', '580', '539', '17.4', '1.17', '1.50']],
    'Table 3. Activity-cliff tiers.',
    'Tiers are mutually exclusive |ΔpIC50| bands; all require Tanimoto ≥ 0.8 (ECFP4, Morgan r = 2, 2,048 bits). '
    'Pair counts are additive (867 total). Compound counts are NOT additive across tiers (one compound may join '
    'pairs of different tiers); the union of all cliff-participating compounds is 654. % of dataset = compounds in '
    'tier / 3,093. 80 of 94 severe pairs (85.1%) are MMP-confirmed.')

add_table(
    ['Hub', 'InChIKey (skeleton)', 'Class', 'pIC50', 'Severe pairs', '% of 94', 'Series size'],
    [['A1', 'SMADULGDNOCLOP', 'A (series floor)', '5.39', '15', '16.0', '174'],
     ['A2', 'RAVBZQAQTVGKIV', 'A (series floor)', '5.34', '12', '12.8', '174'],
     ['B1', 'UDCDEKJNAMHBFH', 'B (singleton attractor)', '4.30', '12', '12.8', '1'],
     ['B2', 'DVCKJOQIVOGXEI', 'B (singleton attractor)', '4.30', '11', '11.7', '1']],
    'Table 4. Cliff-hub compounds.',
    'Together the four hubs participate in 50 of 94 severe cliff pairs (53.2%). Class B hubs differ by a single '
    'methylene (mutual Tanimoto 0.975).')

# ══════════════════════════════════════════════════════════════════════════════
# SUPPLEMENTARY
# ══════════════════════════════════════════════════════════════════════════════
page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('SUPPLEMENTARY INFORMATION'), size=14, bold=True)
doc.add_paragraph()

add_figure('fig_s01_scaffold_cliff_density.png',
           'Supplementary Figure S1. Scaffold cliff-density ranking — which PAD4 chemotypes produce rugged SAR.',
           '(a) The 11 of 155 Murcko scaffold series (n ≥ 4) that harbor any within-scaffold severe cliff, ranked '
           'by cliff count and annotated with series size and cliff density ρ. (b) Ruggedness is scaffold-intrinsic, '
           'not a sampling artifact: the 174-member rank-1 series (23 cliffs, rugged) and the 102-member rank-2 '
           'series (0 cliffs, smooth, σ = 0.37) are equally densely sampled. Full ranking in Supplementary Table '
           'S-scaf.')
add_figure('fig_s02_assay_enrichment.png', 'Supplementary Figure S2. Assay-class cliff enrichment.',
           'Fisher’s exact test of each assay-mechanism class for over-representation among severe-cliff '
           'compounds. No class is significantly enriched (all P > 0.05); enzymatic classes appear in cliffs '
           'proportionally to their dataset frequency — a useful negative result arguing against an assay-format '
           'origin for the cliffs.')
add_figure('fig_s03_sas_map.png', 'Supplementary Figure S3. Structure–activity similarity (SAS) map.',
           '(a) Hexbin density (log) of all 358,416 related pairs (Tanimoto vs |ΔpIC50|); 94 severe cliffs '
           'highlighted in the upper-right activity-cliff quadrant. The map partitions into smooth/continuous SAR '
           '(high similarity, low ΔpIC50), activity cliffs (high similarity, high ΔpIC50), discontinuous/scaffold-'
           'hop pairs (low similarity, high ΔpIC50) and non-descript pairs (low similarity, low ΔpIC50). '
           '(b) Per-similarity-bin |ΔpIC50| distribution and cliff rate, demonstrating diagonal absence. Quadrant '
           'counts: cliffs 0.026%, smooth SAR 3.34%, discontinuous 0.57%, non-descript 96.06%.')

add_md('**Supplementary Table S-hub.** Hub (n = 4) versus non-hub cliff-compound (n = 95) physicochemical '
       'comparison (Mann–Whitney U). Only pIC50 differs significantly (hub 4.83 vs non-hub 6.74, P = 0.007); '
       'all physicochemical descriptors P > 0.26.', space_after=4)
add_md('**Supplementary Table S-scaf.** Per-scaffold cliff-density ranking (Scaffold rank | Compounds | Cliffs | '
       'Cliff density | σ | mean pIC50); 11 of 155 series (n ≥ 4) harbor any within-scaffold cliff.', space_after=4)
add_md('**Supplementary Table S-SAS.** SAS quadrant distribution of all 358,416 related pairs.', space_after=4)
add_md('**Supplementary Table S5.** Top SALI pairs (deduplicated, 17 rows). Full 94-pair severe-cliff table, '
       'fingerprint-sensitivity table, and complete 3,093-compound list provided as machine-readable files in the '
       'deposited dataset.', space_after=4)

# ── Citations needed ──────────────────────────────────────────────────────────
add_heading('Citations required before submission', 3)
for c in [
    'PAD4 biology / citrullination and NETosis [ref]',
    'Rheumatoid arthritis and anti-citrullinated protein antibodies [ref]',
    'Haloacetamidine covalent inhibitors (Cl-amidine/F-amidine; Causey, Thompson et al.) [ref]',
    'Reversible benzimidazole inhibitors (GSK484/GSK199; Lewis et al.) [ref]',
    'Stumpfe & Bajorath (2012) — activity-cliff / ECFP4 threshold justification',
    'Senger (2009) — activity-cliff definition (Tanimoto ≥ 0.8, ΔpIC50 ≥ 2.0)',
    'Knuckley (2010) or equivalent — PAD4 calcium dependence (JBI-589 assay-condition note)',
    'ChEMBL, BindingDB, PubChem database references; RDKit; t-SNE; UpSet; SALI (Guha & Van Drie) [ref]',
]:
    pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.5)
    pp.paragraph_format.first_line_indent = Cm(-0.3); pp.paragraph_format.space_after = Pt(2)
    set_run_font(pp.add_run('• ' + c), size=9)

doc.save(str(OUT))
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size / 1024:.0f} KB')
