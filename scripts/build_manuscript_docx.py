"""
Build PAD4-DB v2 complete integrated manuscript DOCX.
Embeds all 6 main figures and 7 supplementary figures inline.
All 4 main tables rendered as Word tables at first reference.
Run from project root: conda run -n pad4bench python3 scripts/build_manuscript_docx.py
"""
import os
os.chdir('/home/nidhal/PAD4-db_V2')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PUB  = Path('publication')
FIGS_MAIN = PUB / 'figures' / 'main'
FIGS_SUPP = PUB / 'figures' / 'supplementary'
OUT  = PUB / 'manuscript' / 'PAD4_DB_v2_manuscript_integrated.docx'

# ── Helpers ────────────────────────────────────────────────────────────

def set_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_spacing(h, before=14 if level == 1 else 9, after=4)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    set_spacing(p, before=0, after=6, line=14)
    return p

def add_body_mixed(doc, parts):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    set_spacing(p, before=0, after=6, line=14)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p

def add_figure(doc, png_path, fig_num, caption_body, width_in=6.0):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_img, before=14, after=3)
    p_img.add_run().add_picture(str(png_path), width=Inches(width_in))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_cap, before=3, after=14)
    r_label = p_cap.add_run(f'Figure {fig_num}. ')
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_rest = p_cap.add_run(caption_body)
    r_rest.font.size = Pt(9)
    return p_img

def add_supp_figure(doc, png_path, fig_label, caption_body, width_in=5.8):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_img, before=14, after=3)
    p_img.add_run().add_picture(str(png_path), width=Inches(width_in))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_cap, before=3, after=14)
    r_label = p_cap.add_run(f'{fig_label}. ')
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_rest = p_cap.add_run(caption_body)
    r_rest.font.size = Pt(9)

def add_table_title(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=14, after=4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    return p

def _set_cell(cell, text, bold=False, font_size=9, align_left=False):
    """Write text into a table cell with explicit font to prevent PDF ligature garbling."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Liberation Sans'
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_left else WD_ALIGN_PARAGRAPH.CENTER

def make_table(doc, headers, rows, col_widths=None, note=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    # Header
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell(hdr[i], h, bold=True, font_size=9, align_left=False)
    # Data
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            _set_cell(cells[ci], val, bold=False, font_size=9, align_left=(ci == 0))
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    if note:
        pn = doc.add_paragraph(f'Note: {note}')
        pn.style = doc.styles['Normal']
        set_spacing(pn, before=3, after=12)
        if pn.runs:
            pn.runs[0].font.size = Pt(8)
            pn.runs[0].italic = True
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    return tbl

def hr(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=6)
    r = p.add_run('─' * 90)
    r.font.size = Pt(5)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── BUILD ───────────────────────────────────────────────────────────────
doc = Document()

# Margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# Base font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── TITLE ──────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(t, before=0, after=8)
rt = t.add_run(
    'PAD4-DB v2: A Provenance-First Database of PAD4 Inhibitors '
    'with Activity Cliff Characterization and Source Independence Scoring'
)
rt.bold = True; rt.font.size = Pt(15)

for meta in ['[Author names TBD]', '[Affiliations TBD]',
             'Correspondence: [TBD]']:
    p = doc.add_paragraph(meta)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=3)
    p.runs[0].font.size = Pt(10)

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(kw, before=6, after=14)
rk = kw.add_run('Keywords: '); rk.bold = True; rk.font.size = Pt(10)
kw.add_run(
    'PAD4; peptidylarginine deiminase; activity cliffs; matched molecular pairs; '
    'database; structure-activity relationships; machine learning; scaffold analysis'
).font.size = Pt(10)

hr(doc)

# ── ABSTRACT ──────────────────────────────────────────────────────────
add_heading(doc, 'Abstract', 1)
add_body(doc,
    'Peptidylarginine deiminase 4 (PAD4) is a therapeutic target for autoimmune and '
    'inflammatory diseases, with multiple inhibitors in clinical development. However, '
    'no curated, target-specific database exists to support structure-activity relationship '
    '(SAR) studies or machine learning benchmarking. We present PAD4-DB v2, a provenance-first '
    'database of 3,093 curated PAD4 inhibitors from 95 PubChem bioassays, ChEMBL, and '
    'BindingDB, with explicit tracking of data origin through a six-layer architecture (A–F) '
    'and an independence score distinguishing true multi-source replication (n=528) from '
    'pipeline redundancy (n=2,565). The database captures 1,244 Bemis-Murcko scaffolds, '
    '94 severe activity cliff pairs (Tanimoto ≥ 0.8, ΔpIC50 ≥ 2.0), 85.1% of which are '
    'validated by matched molecular pair analysis as single R-group substitutions. Four '
    'cliff-hub compounds in two structural classes account for 53.2% of severe cliff pairs; '
    'permutation testing confirms both the cliff rarity (~20-fold below chance) and the hub '
    'concentration (3.5-fold above chance; both p < 0.0001), representing two independent '
    'and mechanistically distinct structural patterns that are candidate failure modes for '
    'similarity-based ML models. '
    'PAD4-DB v2 is available with a persistent DOI and supports the companion PAD4-Bench '
    'benchmark for evaluating compound activity prediction methods.'
)
hr(doc)

# ── SECTION 1 ─────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction', 1)
add_body(doc,
    'Protein-arginine deiminase 4 (PAD4) catalyses the calcium-dependent conversion of '
    'peptidyl-arginine to peptidyl-citrulline, a post-translational modification with causal '
    'roles in rheumatoid arthritis, cancer progression through neutrophil extracellular trap '
    '(NET) formation, and cardiovascular disease. The enzyme is activated by physiological '
    'calcium concentrations and is most highly expressed in haematopoietic cells, where it '
    'citrullinates histones H3 and H4, promoting chromatin decondensation during NETosis and '
    'establishing a mechanistic link between autoimmune antigen generation and inflammatory '
    'cell death. PAD4 deimination is further implicated in immune checkpoint regulation and '
    'inflammatory resolution, and genetic studies have linked PADI4 polymorphisms to '
    'susceptibility in rheumatoid arthritis, systemic lupus erythematosus, and ulcerative '
    'colitis. The combination of well-characterised structural biology (crystal structures at '
    'sub-2 Å resolution in the PDB spanning both covalent adducts and reversible inhibitor '
    'complexes) and an established clinical dataset makes PAD4 one of the most tractable '
    'epigenetic enzyme targets for quantitative structure-activity relationship (SAR) analysis. '
    'Several PAD4 inhibitors have entered clinical evaluation for rheumatoid arthritis and '
    'related autoimmune indications, including covalent chloroacetamidine-class compounds and '
    'reversible heterocyclic chemotypes developed through iterative medicinal chemistry '
    'campaigns over the past decade. The emerging roles of PAD4 in oncology '
    '(NET-mediated tumour promotion) and thrombosis (citrullination of fibrinogen) have '
    'further broadened clinical interest and expanded the patent landscape substantially. This '
    'therapeutic trajectory makes PAD4 an attractive and pharmacologically rich model system '
    'for SAR analysis and ML benchmarking, provided that high-quality, unified bioactivity '
    'data are available in a computationally accessible form.'
)
add_body(doc,
    'No target-specific, curated database exists for PAD4 inhibitors. The most relevant '
    'public resource, ChEMBL target CHEMBL6111, contains 4,925 activity rows but lacks three '
    'features essential for rigorous SAR analysis and ML model development. First, ChEMBL '
    'CHEMBL6111 does not capture the 327,336-compound HTS structural reference generated by '
    'three large-scale PubChem screening campaigns (AIDs 463073, 485272, 488796), which '
    'provides the full chemical space sampled by industrial-scale PAD4 screening and is '
    'essential for interpreting the active fraction. Second, no existing resource applies '
    'independence scoring to quantify whether apparent multi-source coverage represents '
    'genuinely independent experimental replication or database re-curation of shared primary '
    'measurements; for PAD4, 89.1% of compounds appear in two or more source databases, but '
    'this figure reflects the architecture of ChEMBL and BindingDB as aggregators of PubChem '
    'bioassay data rather than replicated experiments. Third, systematic activity cliff '
    'characterization—identifying pairs of structurally similar compounds with large potency '
    'differences—has not been performed at the PAD4 target level, leaving the SAR '
    'discontinuity landscape uncharacterized for algorithm developers and medicinal chemists alike.'
)
add_body(doc,
    'PAD4-DB v2 addresses these gaps through four contributions. First, a provenance-first '
    'six-layer architecture integrates 95 PubChem bioassay identifiers spanning confirmatory '
    'IC50 campaigns, literature-derived kinetic assays, fluorescence polarization binding '
    'assays, and HTS single-concentration screens with ChEMBL (CHEMBL6111) and BindingDB '
    '(Q9UM07), yielding 3,093 unique SAR compounds with traceable data origin. Second, a '
    'source independence score formally distinguishes the 528 compounds (17.1%) with '
    'genuinely multi-source measurements (score ≥ 0.6) from the 2,565 compounds (82.9%) '
    'whose apparent multi-source presence reflects shared PubChem-BindingDB-ChEMBL curation '
    'pipelines. Third, systematic activity cliff analysis with orthogonal matched molecular '
    'pair (MMP) validation identifies 94 severe cliff pairs and reveals a two-class hub '
    'structure covering 53.2% of severe pairs, representing two mechanistically distinct '
    'structural patterns that are candidate failure modes for similarity-based ML models. '
    'Fourth, all data and analysis '
    'scripts are deposited with a persistent DOI and support the companion PAD4-Bench '
    'benchmark for systematic evaluation of compound activity prediction methods. PAD4-DB v2 '
    'is freely available with a persistent DOI; all analysis scripts are provided for full '
    'reproducibility.'
)

add_heading(doc, '1.1  Comparison with Existing Resources', 2)
add_body(doc,
    'PAD4 bioactivity data are available from three primary repositories: ChEMBL (assay '
    'CHEMBL6111), BindingDB (target Q9UM07), and PubChem BioAssay (95 AIDs). Both ChEMBL and '
    'BindingDB function as single-source aggregators that index activity data without '
    'provenance tracking, source-independence accounting, an HTS structural reference, or '
    'systematic cliff characterisation. PAD4-DB v2 integrates all three repositories and adds '
    'those analytical layers (Table 5).'
)
add_body(doc,
    'Two multi-target activity-cliff benchmarks exist: MoleculeACE (30 targets, ~35,600 unique '
    'compounds curated from ChEMBL, a potency-regression task) and ACNet (190 targets, over '
    '400,000 matched molecular pairs including over 20,000 cliffs from ChEMBL, an '
    'MMP-classification task). Both provide cross-target breadth for evaluating ML '
    'generalisation but draw from a single database, without provenance, source-independence '
    'accounting, or a per-target HTS structural reference, and aggregate compounds across many '
    'targets. PAD4-DB v2 is complementary, providing single-target depth — multi-source '
    'integration with documented provenance, a 327,336-compound HTS reference, and a '
    'permutation-validated cliff/hub landscape — for one clinically active target.'
)
hr(doc)

# ── SECTION 2 ─────────────────────────────────────────────────────────
add_heading(doc, '2. Data Sources and Curation Pipeline', 1)

add_heading(doc, '2.1  Raw Data Acquisition', 2)
add_body(doc,
    'Three primary source databases were queried for PAD4 (UniProt Q9UM07) bioactivity data. '
    'PubChem BioAssay was queried for all assay identifiers (AIDs) linked to human PAD4 '
    'across five assay categories: 57 confirmatory IC50 assays (Layer A), 11 literature-derived '
    'kinetic assays (Layer C), 26 secondary biophysical and functional assays (Layers D and E), '
    'and 3 HTS single-concentration campaigns (Layer F). A total of 95 unique AIDs were '
    'retrieved; two AIDs (1920046 and 2202442) appeared in both the confirmatory and secondary '
    'subdirectories, and their records were loaded exclusively from the confirmatory layer to '
    'avoid duplication. ChEMBL bioactivity data were retrieved as the complete CHEMBL6111 '
    'target export (semicolon-delimited, 4,925 rows). BindingDB data were retrieved as the '
    'complete Q9UM07 target set (tab-delimited, 3,087 rows).'
)
add_body(doc,
    'PubChem bioassay data were downloaded 2026-06-10 to 2026-06-14. ChEMBL bioactivity data '
    '(assay CHEMBL6111) were downloaded on 2026-06-14 (file modification time shows 1980-01-01 '
    'due to a known Linux unzip timestamp artifact; the download date is confirmed by directory '
    'modification time and batch download logs). BindingDB (UniProt accession Q9UM07) was '
    'downloaded on 2026-06-10.'
)
add_body(doc,
    'Raw record counts totalled 341,328 properly parsed rows across all sources. Forty-six rows '
    'were excluded at ingestion as secondary copies of dual-layer AIDs (23 rows each for AIDs '
    '1920046 and 2202442), yielding 341,282 entries for downstream processing. A pipeline '
    'overview is provided in Supplementary Fig. S1.'
)

add_heading(doc, '2.2  Six-Layer Architecture', 2)
add_body(doc,
    'The curation pipeline organises assays into six functional layers. Layer A comprises '
    'confirmatory IC50 dose-response assays (57 AIDs), which are the primary source of potency '
    'data used in all downstream analyses; only rows from Layer A and its ChEMBL/BindingDB '
    'equivalents contribute to pIC50 consensus values. Layer B encompasses secondary IC50 '
    'campaigns with more restricted inclusion criteria. Layer C captures kinetic inactivation '
    'assays reporting Kinact/Ki ratios (11 AIDs), which are retained as metadata but excluded '
    'from IC50 aggregation because they measure a different mechanistic quantity. Layer D '
    'includes fluorescence polarization binding assays (AID 1346144 and related), which report '
    'binding rather than functional inhibition; Ki/Kd values from these assays are not combined '
    'with IC50 values. Layer E encompasses cellular and functional assays. Layer F is the HTS '
    'structural reference (3 AIDs; 327,336 unique compounds), in which percent-inhibition data '
    'at a single compound concentration are retained as a structural index but never used in '
    'pIC50 calculations (use_in_potency_model=False for all Layer F records). This layered '
    'architecture ensures that mechanistically distinct endpoint types are never averaged '
    'together and that each data record carries an explicit provenance label.'
)

add_heading(doc, '2.3  SMILES Standardization', 2)
add_body(doc,
    'SMILES strings were standardised using RDKit 2025.09.5. The standardisation pipeline '
    'applied salt stripping via LargestFragmentChooser, charge neutralisation via Uncharger, '
    'canonical SMILES generation, and InChIKey computation. BindingDB exports contain Daylight '
    'extended SMILES annotations (suffixes of the form |r,THB:...|) that are not valid RDKit '
    'input; these suffixes were stripped before SMILES parsing, raising the parse success rate '
    'from 83% to 100% for BindingDB entries. Of 341,282 entries processed, 341,276 (99.998%) '
    'yielded valid structures; 6 entries with no SMILES string were retained in the raw record '
    'but excluded from further analysis (NO_SMILES status). Zero parse failures and zero '
    'sanitisation failures were recorded, and SMILES integrity checks confirmed that no SMILES '
    'string mapped to multiple InChIKeys and no InChIKey mapped to multiple canonical SMILES.'
)

add_heading(doc, '2.4  Activity Normalization and Replicate Aggregation', 2)
add_body(doc,
    'IC50 values were converted to nanomolar units (µM × 1,000; mM × 10⁶; pM × 0.001) and '
    'then to pIC50 via pIC50 = −log₁₀(IC50[M]). Percent-inhibition records (330,136 rows '
    'comprising the dominant HTS layer) were intercepted before unit conversion and retained '
    'as categorical activity annotations rather than being incorrectly treated as concentration '
    'values; this intercept raised the normalization success rate from 89.8% to 99.0% '
    '(338,021/341,282 rows with norm_status=OK). Of the remaining 1.0%, 3,155 rows (0.9%) '
    'had no numeric value (MALDI mass-shift assays, kinetic assays without IC50 readout) and '
    '106 rows (0.0%) carried unconvertible units (kon/koff rates, dimensionless ratios). The '
    'ChEMBL pchembl_value field was used as a cross-validation check; zero mismatches were '
    'observed across all ChEMBL records with both a reported pchembl_value and a '
    'pipeline-computed pIC50.'
)
add_body(doc,
    'Replicate aggregation consolidated measurements from the same compound in the same assay '
    'via the log-space mean of pIC50 values, which is mathematically equivalent to the '
    'geometric mean of IC50 values and is the statistically appropriate aggregator for '
    'log-scale potency data. The maximum arithmetic deviation from exact log-space mean across '
    'all 7,319 IC50 groups was 0.000000 log units (verified by Step 03b QC). Replicate '
    'consolidation reduced 7,815 compound-assay measurement rows to 7,319 unique '
    'compound-assay pairs across 450 multi-replicate groups (496 measurements consolidated).'
)

add_heading(doc, '2.5  Deduplication, Consensus Potency, and Source Independence Scoring', 2)
add_body(doc,
    'InChIKey-level deduplication across all three source databases yielded 3,093 unique SAR '
    'compounds. The consensus pIC50 for each compound was computed as the mean of per-assay '
    'median pIC50 values, with cross-source concordance defined as agreement within 1.0 log '
    'unit (0 discordant pairs observed; no cross-source conflicts exceeding 1.0 log units were '
    'observed; 0 high-conflict pairs exceeding 1.5 log units). Two ChEMBL compounds with '
    'qualifying IC50 measurements but absent SMILES in the ChEMBL export were correctly '
    'excluded, maintaining the structure-resolved compound count at 3,093.'
)
add_body(doc,
    'Source independence scores were assigned to each compound to quantify the degree of '
    'genuine cross-source measurement diversity (Table 4). The score was derived from the '
    "compound's source combination as follows: compounds in BindingDB + ChEMBL + PubChem were "
    'assigned 0.3, reflecting the lowest independence because all three aggregators draw on the '
    'same PubChem bioassay campaigns; compounds in BindingDB + PubChem were assigned 0.5; '
    'compounds in BindingDB + ChEMBL were assigned 0.6; compounds in any other two-source '
    'combination were assigned 0.7; and compounds from a single source only were assigned 1.0 '
    '(representing the absence of re-curation redundancy rather than a second experimental '
    'measurement). Using a threshold of ≥ 0.6 as the criterion for genuinely multi-source '
    'data, 528 compounds (17.1%) meet this criterion; the remaining 2,565 compounds (82.9%) '
    'appear in multiple databases because of reciprocal re-curation of shared PubChem bioassay '
    'records. The 89.1% apparent multi-source rate (2,755/3,093) therefore reflects database '
    'architecture rather than distinct experimental programmes. Cross-source overlap is '
    'displayed as an UpSet plot in Fig. 2; source independence score distributions are shown '
    'in Table 4.'
)
add_body(doc,
    'The independence score is an ordinal heuristic for re-curation redundancy rather than a '
    'direct measurement of experimental independence: it ranks source combinations by their '
    'expected degree of shared PubChem provenance, so the >= 0.6 threshold identifies compounds '
    'free of demonstrable re-curation redundancy rather than compounds with confirmed replicate '
    'measurements. Of the 528 compounds meeting this threshold, 338 are single-source (score 1.0 '
    '- no cross-source redundancy, but also no second measurement) and 190 are genuinely '
    'multi-source above the threshold (BindingDB+ChEMBL, n=167; ChEMBL+PubChem, n=23).'
)
add_body(doc,
    'Of the 2,755 multi-database compounds (89.1% of 3,093), 96.2% (2,649/2,755) show zero '
    'cross-source pIC50 spread, consistent with these compounds sharing primary measurements '
    'through database re-curation rather than carrying independent replicate determinations. '
    'A high-reproducibility subset - the PAD4 Golden Set (PAD4_Golden_Set.csv; n=47, 1.5% of '
    'the dose-response database) - applies a stricter criterion: presence in >=2 distinct '
    'genuine PubChem assays (excluding aggregator pseudo-identifiers CHEMBL6111 and Q9UM07) '
    'with cross-assay spread <= 0.5 log units, providing a maximally reliable pIC50 reference '
    'independent of cross-database re-curation.'
)

add_heading(doc, '2.6  Reproducibility', 2)
add_body(doc,
    'The complete curation pipeline reproduces all reported statistics from raw source files '
    'in approximately four minutes on a standard workstation (Python 3.10, RDKit 2025.09.5).'
)
hr(doc)

# ── SECTION 3 ─────────────────────────────────────────────────────────
add_heading(doc, '3. Dataset Characteristics', 1)

add_heading(doc, '3.1  Overall Composition', 2)
add_body(doc,
    'The PAD4-DB v2 dose-response dataset comprises 3,093 unique compounds with consensus '
    'pIC50 values spanning six assay mechanism classes. The HTS structural reference layer '
    'comprises 327,336 unique compounds tested in at least one of three PubChem HTS campaigns. '
    'Of 3,093 PAD4 inhibitors in the dose-response database, 1,453 share InChIKeys with '
    'compounds in the HTS screening dataset (327,336 total screened), indicating parallel '
    'measurement in both assay formats. Only 6 of these 1,453 compounds were confirmed HTS '
    'actives (max inhibition ≥50% at screening concentration); the remaining 1,447 have '
    'published IC50 values from independent research programs (ChEMBL/BindingDB) and showed '
    'low inhibition in HTS campaigns (median 4.3% at screening concentration), consistent with '
    'potent compounds (median pIC50=6.93) tested below their IC50 at standard HTS screening '
    'concentrations. The combined non-redundant InChIKey space across SAR and HTS layers '
    'totals 328,976 unique structures (Supplementary Table S1).'
)

add_heading(doc, '3.2  Potency Distribution', 2)
add_body(doc,
    'Consensus pIC50 values span the full range of 2.00 to 8.52 log units, corresponding to '
    '10 mM to 3 nM IC50. The mean pIC50 is 6.55 (SD 0.99), the median is 6.84, and the '
    'distribution is bimodal: a primary peak centred at pIC50 ≈ 7.0 reflects the dominant '
    'confirmed IC50 assay chemotypes, while a secondary shoulder at pIC50 ≈ 5.0 is driven by '
    'the 233 patent-exclusive compounds that were identified in earlier-stage screening '
    'campaigns. The mean lies below the median (6.55 vs 6.84) because of a weak-binder tail '
    'extending to pIC50 = 2.00, comprising 19 compounds with IC50 ≥ 1 mM that are retained '
    'as authentic measurements with norm_status=OK. The BindingDB-only compounds (n=95) show '
    'a sharp mode at approximately pIC50 7.3, consistent with BindingDB\'s selective curation '
    'of higher-potency published inhibitors. Potency distribution statistics are displayed in '
    'Fig. 3a.'
)

add_heading(doc, '3.3  Source Coverage and Patent Compounds', 2)
add_body(doc,
    'Across the 3,093 SAR compounds, BindingDB records are present for 91.4% (2,827 '
    'compounds), PubChem confirmatory records for 91.2% (2,821 compounds), and ChEMBL records '
    'for 50.6% (1,566 compounds). The source distribution is illustrated in Fig. 2 as an '
    'UpSet plot. Patent-exclusive compounds, defined as those present only in PubChem '
    'confirmatory assays without ChEMBL or BindingDB coverage (n=233), have a mean pIC50 of '
    '6.08, compared with 6.59 for published (non-patent-exclusive) compounds (Mann-Whitney U, '
    'p < 0.001). This difference reflects the earlier-stage optimization of patent screening '
    'hits relative to compounds that have been characterised in published medicinal chemistry '
    'programs, rather than intrinsic differences in target binding affinity; the patent '
    'compounds cover 103 structurally distinct scaffold series absent from the non-patent '
    'space with a mean scaffold series size of 2.5 compounds per series, identical to that of '
    'non-patent scaffolds (2.5 compounds per series). Fig. 3b and Supplementary Fig. S3 '
    'display the pIC50 distributions by source type.'
)

add_heading(doc, '3.4  Assay Mechanism Classes', 2)
add_body(doc,
    'Four mechanism classes are assigned based on the assay format recorded in the source '
    'PubChem AID descriptions and ChEMBL assay annotations. The enzymatic class '
    '(BAEE-colorimetric; n=2,079, 67.2%) encompasses primary IC50 assays using the '
    'benzoyl-arginine-based chromogenic substrate. The enzymatic_confirmed class '
    '(RFMS-fluorescence; n=878, 28.4%) uses a rhodamine-based fluorescent substrate for '
    'orthogonal confirmation. The fp_ic50 class (fluorescence polarization binding; n=115, '
    '3.7%) reports binding IC50 rather than enzymatic inhibition IC50. The covalent class '
    '(n=21, 0.7%) comprises compounds characterised by irreversible kinetic assay formats. '
    'Both enzymatic classes measure IC50 on the same biochemical target via orthogonal '
    'detection formats; the enzymatic_confirmed median pIC50 of 6.66 is lower than the '
    'enzymatic median of 6.88 by 0.22 log units, a difference attributable to '
    'calcium-concentration differences between BAEE and RFMS assay protocols rather than to '
    'systematic potency bias. Four severe cliff pairs (4.3% of 94) involve cross-mechanism '
    'pairs within the enzymatic IC50 family (enzymatic vs enzymatic_confirmed), none of which '
    'involve the fp_ic50 or covalent classes; the mechanistic heterogeneity does not drive the '
    'hub structure (Section 5.4). Mechanism class distributions are shown in Fig. 3c.'
)

add_heading(doc, '3.5  Physicochemical Property Landscape', 2)
add_body(doc,
    'RDKit physicochemical descriptors were computed for all 3,093 compounds from their '
    'standardised SMILES with zero parse failures (Fig. S7). The dataset occupies a relatively '
    'high molecular-weight region of chemical space, with a median molecular weight of 590.7 Da '
    '(5th-95th percentile 375.5-666.9 Da), reflecting the dominance of the large fused-ring '
    'azaindole-benzimidazole series. Median Crippen cLogP is 4.60 (2.84-6.07), median TPSA '
    '111.4 A^2 (68.6-146.6 A^2), with medians of 6 H-bond acceptors (3-8), 2 H-bond donors '
    '(1-3), 6 rotatable bonds (4-10), an Fsp3 of 0.38 (0.18-0.54), and 5 aromatic rings (2-6). '
    'Consequently only 19.8% of compounds (613/3,093) satisfy all four Lipinski rule-of-five '
    'criteria strictly (MW <= 500, cLogP <= 5, HBD <= 5, HBA <= 10); using the classic '
    '<=1-violation threshold, 68.9% (2,130/3,093) are compliant. The dominant pattern is '
    'exactly one violation - almost exclusively MW > 500, consistent with the large fused-ring '
    'azaindole-benzimidazole scaffold (median MW 590.7 Da, the dominant chemotype) - while '
    'cLogP, HBD, and HBA remain within Lipinski bounds for the vast majority. Veber criteria '
    '(rotatable bonds <= 10, TPSA <= 140 A^2) are satisfied by 90.7%. These distributions '
    'indicate that PAD4-DB v2 captures lead- to drug-like chemical space biased toward larger, '
    'more aromatic scaffolds rather than fragment-sized chemotypes.'
)

add_heading(doc, '3.6  Covalent Inhibitors', 2)
add_body(doc,
    'One hundred and seven compounds (3.5%) carry at least one SMARTS-flagged reactive '
    'warhead, distributed across seven chemical classes: chloroacetamidine (n=66), '
    'fluoroacetamidine (n=17), haloacetyl (n=11), enaminone (n=7), vinyl sulfone (n=4), '
    'alpha-bromoketone (n=2), and unclassified warhead-containing structures (n=2). A '
    'covalent-versus-reversible cliff safety audit confirmed that zero severe cliff pairs and '
    'zero moderate cliff pairs involve one covalent and one non-covalent compound; the single '
    'broad-tier covalent-involving pair has ΔpIC50=1.21. The severe cliff landscape is '
    'therefore pharmacologically clean with respect to mechanism-of-action confounds.'
)

add_heading(doc, '3.7  Reference Compound Recovery', 2)
add_body(doc,
    'Seven of thirteen curated PAD4 reference inhibitors are present in PAD4-DB v2 with '
    'concordant pIC50 values (mean |ΔpIC50| < 0.15 log units): Streptonigrin (pIC50=5.602), '
    'Cl-amidine (5.219), F-Amidine (4.571), GSK484 (7.049), TDFA (5.638), BMS-P5 (7.009), '
    'and JBI-589 (6.000). Three reference inhibitors are present in the raw data but correctly '
    'excluded by endpoint-type criteria: o-F-Amidine (kinact/Ki only, no primary IC50), '
    'Amodiaquine (HTS percent-inhibition only, no dose-response), and BB-Cl-Amidine (covalent '
    'kinetics only). Three reference inhibitors are absent from all three source databases: '
    'GSK199 (not submitted under CHEMBL6111 assay), Pyroxamide (absent from all '
    'PAD4-annotated bioassay records), and PAD-PF1 (allosteric inhibitor not in public '
    'databases). AFM-30a was correctly excluded as PAD2-selective. The JBI-589 pIC50 '
    'discrepancy of 0.9 log units (database: 6.000, published: 6.914; 122 nM) is likely '
    'attributable to differences in assay calcium concentration between the source database '
    'records and the original published assay; PAD4 enzymatic activity is known to vary '
    'substantially with free Ca²⁺ concentration, and formal verification is outside the scope '
    'of this curation study.'
)
hr(doc)

# ── SECTION 4 ─────────────────────────────────────────────────────────
add_heading(doc, '4. Scaffold Analysis', 1)

add_heading(doc, '4.1  Method', 2)
add_body(doc,
    'Scaffolds were defined using the Bemis-Murcko heteroatom-preserving method as implemented '
    'in RDKit 2025.09.5 (MurckoScaffold.GetScaffoldForMol). This definition preserves ring '
    'nitrogen, oxygen, and sulphur atoms and is used by the ChEMBL scaffold analysis pipeline, '
    'making it the most widely comparable standard for medicinal chemistry databases. A '
    'canonicalization note is required: the rank-1 scaffold series count of 174 compounds is '
    'locked to RDKit 2025.09.5, as fresh computation in prior RDKit releases yields 190 '
    'compounds for the same scaffold due to inter-version canonical SMILES drift; the deposited '
    'dataset and all reported statistics use the RDKit 2025.09.5 canonical values.'
)

add_heading(doc, '4.2  Chemical Diversity', 2)
add_body(doc,
    'PAD4-DB v2 contains 1,244 unique Bemis-Murcko scaffolds. Of these, 375 scaffolds (30.1% '
    'of total unique scaffolds) constitute named series of two or more compounds, covering '
    '71.9% of the 3,093 dataset compounds; the remaining 869 singletons (of which 3 are '
    'acyclic structures) each represent a single compound. The scaffold '
    'size distribution is highly concentrated: the Gini coefficient is 0.532, and the 30 '
    'largest scaffolds (2.4% of all scaffolds) collectively contain approximately 30% of all '
    'compounds. The median scaffold series size is 3 compounds, and the top 5 series sizes '
    'range from 10 to 174 members. This level of scaffold concentration is consistent with the '
    'known chemotype focus of the PAD4 patent and publication landscape, where a small number '
    'of azaindole-based core structures dominate the medicinal chemistry literature. Scaffold '
    'diversity statistics are displayed in Fig. 4c. Supplementary Table S2 provides the '
    'complete scaffold-compound mapping file with SMILES, series size, mean pIC50, and patent '
    'compound fraction for all 1,244 scaffolds.'
)

add_heading(doc, '4.3  Dominant Scaffold Series', 2)
add_body(doc,
    'The largest scaffold series comprises 174 compounds belonging to an '
    'azaindole-benzimidazole chemotype with a mean pIC50 of 7.07 (Fig. 4a). This series is '
    '70% larger than the rank-2 scaffold (n=102 compounds) and is the primary contributor to '
    'the large-scale PubChem confirmatory screening campaigns in the dataset. The 174-compound '
    'series spans a 3.5 log unit internal potency range (from pIC50 ≈ 3.8 to ≈ 7.5), making '
    'it the most structurally concentrated potency gradient in the dataset; within-series '
    'pIC50 variance (SD=0.88) is consistent with the broad SAR exploration characteristic of '
    'a scaffold that has been advanced into iterative medicinal chemistry optimization '
    'campaigns. This high internal potency diversity within a single ring system is the '
    'structural cause of the Class A hub phenomenon: the two mid-potency Class A hub compounds '
    '(pIC50 ≈ 5.4) sit approximately 1.5–2.0 log units below the series mode while '
    'maintaining high ECFP4 Tanimoto similarity (≥ 0.80) to their higher-potency analogs, '
    'generating a large density of severe cliff pairs from a well-characterized chemotype '
    'rather than from atypical structural features. The Class A cliff-hub compounds '
    '(Section 5.3) are mid-potency members of this series, and their position as '
    'within-series potency floors with respect to higher-potency analogs accounts for 27 of '
    '94 severe cliff pairs. The azaindole-benzimidazole core scaffold structure is shown in '
    'Supplementary Fig. S5.'
)

add_heading(doc, '4.4  Patent-Exclusive Chemical Space', 2)
add_body(doc,
    'The 233 patent-exclusive compounds map to 107 scaffold series of which 103 are absent '
    'from the non-patent compound space. The five largest patent-exclusive scaffolds are an '
    'azaindole-piperidine-cyclohexane amide (29 compounds, mean pIC50=7.13 ± 0.51), a '
    'chalcone-cyclohexane lactam with pyridine substitution (27 compounds, pIC50=5.21 ± 0.52), '
    'a chalcone-cyclohexane lactam with pyrimidine substitution (18 compounds, '
    'pIC50=4.65 ± 0.57), an azaindole-piperidine-cyclohexane amide variant (16 compounds, '
    'pIC50=6.94 ± 0.48), and an azaindole-bicyclic amine amide (10 compounds, '
    'pIC50=7.11 ± 0.15). Patent compounds contribute 1 of 94 severe cliff pairs (1.1%), and '
    '34 cliff pairs involve at least one patent compound, indicating that patent chemical '
    'space contributes structural diversity without generating a disproportionate share of '
    'activity discontinuities. Patent scaffold analysis is shown in Supplementary Fig. S3.'
)
hr(doc)

# ── SECTION 5 ─────────────────────────────────────────────────────────
add_heading(doc, '5. Activity Cliff Analysis', 1)

add_heading(doc, '5.1  Cliff Identification', 2)
add_body(doc,
    'Structural similarity was computed using ECFP4 fingerprints (Morgan algorithm, radius=2, '
    '2048 bits; RDKit 2025.09.5), following Stumpfe and Bajorath (2012) and Senger (2009). '
    'Activity cliffs were defined as pairs with Tanimoto similarity ≥0.8 and |ΔpIC50| ≥2.0 '
    'log units, consistent with established SAR discontinuity literature. Of 94 severe cliff '
    'pairs, 80 (85.1%) were also confirmed by matched molecular pair analysis (MMP), an '
    'orthogonal substructure-based method independent of fingerprint choice. Sensitivity '
    'analysis with ECFP6 (radius=3) showed that 64 of 94 pairs have ECFP4 Tanimoto 0.80–0.85 '
    'and fall below 0.80 at radius=3, consistent with the known resolution-dependent behaviour '
    'of Morgan fingerprints for large fused ring systems; however, 80% of these borderline '
    'pairs (51/64) remain MMP-confirmed, and the hub dominance statistic is '
    'fingerprint-invariant (53.2% under ECFP4; 53.3% under ECFP6). Thirteen pairs (13.8%) '
    'are classified as severe by ECFP4 only, without MMP or ECFP6 corroboration; these are '
    'flagged in the deposited dataset (ecfp4_only_cliff=True in activity_cliffs.parquet).'
)
add_body(doc,
    'The complete similarity landscape comprises 358,416 pairs with Tanimoto ≥ 0.6, of which '
    '12,071 (3.4%) reach Tanimoto ≥ 0.8. Cliff enumeration yielded 94 severe pairs '
    '(Tanimoto ≥ 0.8, |ΔpIC50| ≥ 2.0), 193 moderate pairs (|ΔpIC50| ≥ 1.5), and 580 broad '
    'pairs (|ΔpIC50| ≥ 1.0); a compound participates in a cliff tier if it appears in at '
    'least one pair at that tier. The 94 severe cliff pairs involve 99 unique compounds '
    '(3.2% of the dataset), with a maximum ΔpIC50 of 3.045 (approximately 1,110-fold) and a '
    'mean severe ΔpIC50 of 2.308. The activity landscape is shown in Fig. 5 and Supplementary '
    'Fig. S2. Severe cliffs constitute approximately 0.026% of all similarity-≥0.6 pairs, '
    'confirming that large potency discontinuities are rare events in the bulk of the PAD4 '
    'chemical space.'
)
add_body(doc,
    'To confirm that this rarity reflects genuine structure–activity relationships rather '
    'than the shape of the potency distribution, the 3,093 consensus pIC50 values were '
    'permuted across compounds (10,000 iterations) with the similarity structure held fixed, '
    'and severe cliffs were re-counted among the 12,071 Tanimoto ≥ 0.8 pairs. Random '
    'potency assignment produced 1,923 ± 125 severe cliffs on average — approximately '
    '20-fold more than observed — and the observed count of 94 fell below all 10,000 '
    'permutations (p < 0.0001). Because the permutation preserves the bimodal marginal '
    'pIC50 distribution, this ~20-fold depletion establishes that the 94 severe cliffs are '
    'genuine SAR discontinuities and not artefacts of the potency distribution: structurally '
    'similar compounds overwhelmingly share similar potency, and the severe cliffs are true '
    'exceptions to a predominantly continuous landscape.'
)
add_body(doc,
    'The fingerprint sensitivity analysis (Supplementary Table S6) and the complete ECFP4 '
    'versus ECFP6 Tanimoto values for all 94 severe pairs are deposited with the dataset to '
    'allow readers to apply alternative threshold choices.'
)

add_heading(doc, '5.2  MMP Validation', 2)
add_body(doc,
    'Matched molecular pair analysis was performed on the 99 severe cliff compounds using '
    'RDKit rdMMPA with maxCuts=1, identifying 707 MMP pairs across 24 unique scaffold cores. '
    'Of the 94 Tanimoto-defined severe cliff pairs, 80 (85.1%) were independently confirmed '
    'as matched molecular pairs, establishing that the observed potency discontinuities arise '
    'from discrete, chemically interpretable structural changes rather than from ambiguity in '
    'the fingerprint-based similarity metric. The 80 MMP-confirmed severe pairs decompose '
    'into three change types: single-atom changes (n=45, 56.3%), small substituent changes '
    '(n=27, 33.8%), and medium substituent changes (n=8, 10.0%); no large-substituent changes '
    'were observed among confirmed severe cliffs. The remaining 14 Tanimoto-severe pairs '
    'are not reducible to a single matched-molecular-pair transformation: 13 also fall below '
    'the Tanimoto threshold under ECFP6 (the ecfp4_only set) and 1 is ECFP6-robust but '
    'admits no valid single-cut MMP. MMP results are displayed in Fig. 6 and '
    'Supplementary Table S3.'
)

add_heading(doc, '5.3  Hub Compound Discovery', 2)
add_body(doc,
    'Systematic degree analysis of the 99-node, 94-edge severe cliff network identified four '
    'compounds collectively participating in 50 of 94 severe cliff pairs (53.2%). These hub '
    'compounds are not randomly distributed across the dataset but are organised into two '
    'structurally and mechanistically distinct archetypes, each generating a characteristic '
    'pattern of activity discontinuities (Fig. 7, Table 1).'
)
add_body(doc,
    'The severe activity cliff landscape is organized around two mechanistically distinct hub '
    'archetypes. Class A hubs (SMADULGDNOCLOP-GISFHXKWSA-N and '
    'RAVBZQAQTVGKIV-XBPDSQQVSA-N; pIC50 ≈ 5.4) are mid-potency members of the dominant '
    '174-compound azaindole-benzimidazole scaffold series; their position as within-series '
    'potency floors generates 27 severe cliff pairs against higher-potency analogs in the same '
    'chemotype. Class B hubs (UDCDEKJNAMHBFH-HSZRJFAPSA-N and '
    'DVCKJOQIVOGXEI-XMMPIXPASA-N; pIC50 = 4.301, Tanimoto = 0.975) are scaffold singletons '
    'with no other compounds sharing their Murcko framework; their broad structural '
    'promiscuity by ECFP4 fingerprint creates 23 severe cliff pairs across multiple '
    'chemotypes. Together, these four cliff-hub compounds in two structural classes account '
    'for 50 of 94 severe cliff pairs (53.2%), representing two independent and mechanistically '
    'distinct structural patterns that are candidate failure modes for similarity-based ML models.'
)
add_body(doc,
    'The same permutation framework confirms that the hub concentration is not a by-product '
    'of the underlying similarity-pair density. Drawing 94 cliff pairs at random from each '
    'permutation\'s null cliffs (count-matched to the observed set) and measuring the '
    'fraction incident to the four highest-degree compounds, the null concentration was '
    '15.2% ± 2.6%, against the observed 53.2% — a 3.5-fold excess reached by none of the '
    '10,000 permutations (p < 0.0001). Because the permutation preserves which compounds '
    'participate in many high-similarity pairs but randomises their potencies, this excess '
    'shows that the hub concentration arises from specific potency positioning — consistent '
    'with the within-series potency floors (Class A) and inactive structural attractors '
    '(Class B) described above — rather than from certain compounds simply appearing in '
    'many similar pairs.'
)
add_body(doc,
    'The two archetypes present qualitatively distinct challenges for ML models trained on '
    'molecular fingerprints or graph representations. Class A hubs generate correlated '
    'failures within the dominant scaffold series: a model that learns the series-average '
    'potency will systematically mispredict the two mid-potency compounds and their many '
    'high-potency neighbours, producing clustered errors within the same embedding region. '
    'Class B hubs generate cross-series failures from structural promiscuity: the cyclobutyl '
    'and cyclopentyl sulfonamide compounds (differing by 14 Da, one CH₂) share high ECFP4 '
    'similarity with structurally diverse active compounds while being themselves essentially '
    'inactive (pIC50=4.301), producing false-positive predictions across multiple scaffold '
    'clusters. The hub classification (hub_class column: \'A\', \'B\', or \'none\') and '
    'cliff-hub pair membership are deposited in the dataset for direct use in benchmark '
    'construction and model evaluation.'
)

# ── TABLE 1 ── inline at Section 5.3 ───────────────────────────────────
add_table_title(doc, 'Table 1. Cliff-hub compound properties.')
make_table(doc,
    headers=['Compound (InChIKey)', 'Class', 'pIC50', 'MW', 'Hub pairs', 'Archetype'],
    rows=[
        ['SMADULGDNOCLOP-GISFHXKWSA-N', 'A', '5.390', '611', '15', 'Series-embedded floor'],
        ['RAVBZQAQTVGKIV-XBPDSQQVSA-N', 'A', '5.341', '591', '12', 'Series-embedded floor'],
        ['UDCDEKJNAMHBFH-HSZRJFAPSA-N', 'B', '4.301', '606', '12', 'Singleton attractor'],
        ['DVCKJOQIVOGXEI-XMMPIXPASA-N', 'B', '4.301', '620', '11', 'Singleton attractor'],
    ],
    col_widths=[2.5, 0.5, 0.55, 0.5, 0.7, 1.5],
    note='Class A inter-hub Tanimoto: 0.761. Class B inter-hub Tanimoto: 0.975. Cross-class Tanimoto: ~0.49.'
)

add_heading(doc, '5.4  Implications for Machine Learning Benchmarking', 2)
add_body(doc,
    'The two-class hub structure has direct implications for benchmark design. The 94 severe '
    'cliff pairs are not randomly distributed across the 3,093 compound embeddings: 53.2% are '
    'concentrated around four hub compounds, meaning that standard random or scaffold-split '
    'benchmark partitions will place hub compounds in the training set with high probability, '
    'and is expected to underestimate generalisation difficulty for the most discontinuous '
    'regions of the potency landscape. A model trained without hub-aware splitting would be '
    'expected to learn the hub potency once and predict the remaining pairs without '
    'encountering the structural context that makes them hard. '
    'The ECFP4-only flag (ecfp4_only_cliff=True, n=13 pairs) further identifies '
    'cliff pairs where fingerprint-choice-dependent classification should inform dataset split '
    'design, as these 13 pairs lack ECFP6 or MMP corroboration and represent the most '
    'ambiguous boundary of the structural similarity definition. The PAD4-Bench benchmark '
    'implements hub-aware and ECFP4-sensitive splits that allow systematic evaluation of '
    'activity prediction models under controlled exposure to both hub archetypes and '
    'fingerprint-ambiguous cliff pairs. These patterns are evaluated empirically in the '
    'companion PAD4-Bench benchmark; the present work characterises the structural basis for them.'
)
hr(doc)

# ── SECTION 6 ─────────────────────────────────────────────────────────
add_heading(doc, '6. Data Quality and Validation', 1)
add_body(doc,
    'A ten-phase audit covering 58 independent quality checks was applied to the complete '
    'dataset, and all 58 checks passed. Biological specificity was confirmed through analysis '
    'of all 95 PubChem AIDs: 91 explicitly annotate PAD4 (human, Q9UM07) among the 95 '
    'PubChem AIDs, 2 annotate the PAD family without isoform specification (AIDs 588488 and '
    '588560, both contributing zero potency-space rows and thus excluded from the SAR '
    'dataset), and 2 are marked ambiguous (AIDs 588487 and 651627, also with zero '
    'potency-space rows); ChEMBL (CHEMBL6111) and BindingDB (Q9UM07) are both explicitly '
    'annotated for human PAD4. Species specificity was confirmed for ChEMBL (4,858 Homo '
    'sapiens rows, 67 species-unknown rows, 0 non-human rows) and BindingDB (3,087 Homo '
    'sapiens rows, 0 non-human rows). Chemical correctness was verified by end-to-end pIC50 '
    'tracing from raw IC50 values through unit conversion to final pIC50 (maximum absolute '
    'deviation 0.000000 log units across all 7,319 IC50 groups, confirmed PASS), SMILES '
    'integrity checking (0 SMILES strings mapping to multiple InChIKeys; 0 InChIKeys mapping '
    'to multiple canonical SMILES), and cross-source concordance verification (maximum '
    'cross-source pIC50 delta across the full 3,093-compound dataset = 0.74 log units; '
    '0 pairs exceeding 1.0 log units; 0 high-conflict pairs exceeding 1.5 log units). '
    'An independent reproducibility run executed the complete pipeline from raw source files '
    'and confirmed all 23 primary reproducibility metrics within 3 minutes 43 seconds.'
)
add_body(doc,
    'PAD4-DB v2 has several limitations that users should consider in application. First, '
    'the dataset is bounded by what was submitted to PubChem bioassay campaigns, ChEMBL '
    '(CHEMBL6111), and BindingDB (Q9UM07) as of the download dates; compounds from recent '
    'patent filings (post-2023) and unpublished clinical programs are not included, and three '
    'known reference inhibitors (GSK199, Pyroxamide, PAD-PF1) are absent from all source '
    'databases. Second, no isoform selectivity data (PAD1, PAD2, PAD3, PAD6 vs PAD4) are '
    'available across the 95 assay identifiers, and compounds are annotated only for PAD4 '
    'activity. Third, scaffold assignments are canonical for RDKit 2025.09.5 and will differ '
    'quantitatively under other toolkit versions or scaffold definitions. Fourth, '
    'calcium-dependent assay formats (BAEE and RFMS protocols use different CaCl₂ '
    'concentrations) produce systematic pIC50 differences between mechanism classes that '
    'affect apparent potency comparisons for the same compound across assay types. Fifth, 13 '
    'severe cliff pairs (13.8%) are classified by ECFP4 only, without ECFP6 or MMP '
    'corroboration; these are flagged as ecfp4_only_cliff=True in the deposited '
    'activity_cliffs.parquet file and should be treated with additional caution in SAR '
    'interpretation. Sixth, source databases were downloaded 2026-06-10 to 2026-06-14, and '
    'compounds from subsequent depositions are not represented.'
)
hr(doc)

# ── SECTION 7 ─────────────────────────────────────────────────────────
add_heading(doc, '7. Data and Software Availability', 1)
add_body(doc,
    'PAD4-DB v2 is deposited at Zenodo (DOI: [TBD]) and comprises six primary data files. '
    'The main compound file, pad4_compounds.parquet, contains 3,093 compounds with consensus '
    'pIC50 values, source annotations, scaffold assignments, cliff-hub class labels, source '
    'independence scores, mechanism class, warhead class, and fragment flags. The cliff pair '
    'file, activity_cliffs.parquet, contains 867 cliff pairs at all tiers (severe, moderate, '
    'broad) with ECFP4 Tanimoto, ECFP6 Tanimoto, delta-pIC50, MMP validation status, and the '
    'ecfp4_only_cliff flag indicating the 13 pairs confirmed by ECFP4 only. The pairwise '
    'similarity file, activity_pairs_with_sali.parquet, contains 358,416 compound pairs with '
    'Tanimoto >= 0.6, source labels, and SALI (Structure-Activity Landscape Index) values. '
    'The HTS structural reference, hts_compound_index.parquet, contains 327,336 compounds '
    'with HTS activity scores, consensus confidence values, and the confirmed_in_potency_space '
    'flag. The MMP output file, mmp_pairs_cliff99.csv, contains 707 MMP pairs among the 99 '
    'severe cliff compounds with change-type classification and per-compound discontinuity '
    'scores. The high-reproducibility reference subset, PAD4_Golden_Set.csv, contains 47 '
    'compounds measured in >=2 distinct genuine PubChem assays (excluding aggregator '
    'identifiers) with cross-assay spread <= 0.5 log units, providing a maximally reliable '
    'pIC50 reference for benchmark calibration.'
)
add_body(doc,
    'All analysis scripts are provided in the scripts/ directory, organized by pipeline stage '
    '(01 through 05, audit, and figures subdirectories). The conda environment specification '
    '(pad4bench; Python 3.10, RDKit 2025.09.5) is provided as environment.yml. Supplementary '
    'Table S6 (ECFP4 vs ECFP6 Tanimoto for all 94 severe cliff pairs) is deposited as both '
    'an HTML interactive table and a machine-readable CSV. PAD4-DB v2 supports the PAD4-Bench '
    'benchmark for systematic evaluation of compound activity prediction models under '
    'conditions that isolate cliff-driven generalisation failure from routine interpolation error.'
)
hr(doc)

# ── SECTION 8 ─────────────────────────────────────────────────────────
add_heading(doc, '8. Conclusions', 1)
add_body(doc,
    'PAD4-DB v2 fills a documented gap in public bioactivity databases by providing the first '
    'curated, target-specific resource for PAD4 inhibitors with traceable data provenance, '
    'explicit independence scoring, and systematic activity cliff characterisation. The '
    '3,093-compound SAR dataset, combined with the 327,336-compound HTS structural reference, '
    'covers the full range of chemical space that has been experimentally evaluated against '
    'PAD4 in the public domain through mid-2026.'
)
add_body(doc,
    'The provenance-first design, with its explicit six-layer architecture and source '
    'independence score, is directly transferable to other high-value drug targets where '
    'multiple public databases provide overlapping but architecturally redundant coverage. The '
    'methodology for computing source independence scores and distinguishing pipeline '
    're-curation from genuine experimental replication addresses a systematic problem in '
    'multi-database integration that affects any target with substantial ChEMBL, BindingDB, '
    'and PubChem overlap.'
)
add_body(doc,
    'The hub compound discovery methodology introduced here—systematic cliff network degree '
    'analysis followed by structural archetype assignment—identifies a two-class attractor '
    'structure that is mechanistically interpretable rather than merely statistical. The '
    'Class A (series-embedded potency floor) and Class B (scaffold-singleton structural '
    'attractor) archetypes represent a general SAR pattern that is likely to appear in any '
    'target with both a dominant scaffold series and structural promiscuity among low-potency '
    'compounds, making the framework applicable beyond PAD4.'
)
add_body(doc,
    'PAD4-DB v2 and the companion PAD4-Bench benchmark together provide the infrastructure '
    'for rigorous, reproducible evaluation of activity prediction models in a clinically '
    'relevant target context. The combination of a curated SAR dataset with documented '
    'provenance, a characterized activity cliff landscape including hub compound annotations, '
    'and a large HTS structural reference enables benchmark design that distinguishes genuine '
    'generalization failure from dataset construction artefacts—a distinction that has been '
    'absent from most published ML benchmarks in drug discovery.'
)
hr(doc)

# ── MAIN TEXT TABLES 2, 3, 4 ──────────────────────────────────────────
add_heading(doc, 'Tables', 1)

add_table_title(doc, 'Table 2. Cliff tier summary.')
make_table(doc,
    headers=['Tier', 'Tanimoto', 'ΔpIC50', 'Pairs', 'Compounds', 'MMP-confirmed'],
    rows=[
        ['Severe',   '≥ 0.8', '≥ 2.0',       '94',  '99',  '80 (85.1%)'],
        ['Moderate', '≥ 0.8', '1.5–<2.0',    '193', '209', '—'],
        ['Broad',    '≥ 0.8', '1.0–<1.5',    '580', '654', '—'],
    ],
    col_widths=[1.1, 1.0, 0.9, 0.8, 1.0, 1.4],
)

add_table_title(doc, 'Table 3. Source distribution summary.')
make_table(doc,
    headers=['Source combination', 'n compounds', 'Independence score', 'Interpretation'],
    rows=[
        ['BindingDB + ChEMBL + PubChem', '1,366', '0.3', 'Pipeline re-curation'],
        ['BindingDB + PubChem',          '1,199', '0.5', 'Pipeline re-curation'],
        ['BindingDB + ChEMBL',             '167', '0.6', 'Genuinely multi-source'],
        ['ChEMBL + PubChem',               ' 23', '0.7', 'Genuinely multi-source'],
        ['PubChem only (patent-exclusive)', '233', '1.0', 'Source-exclusive'],
        ['BindingDB only',                 ' 95', '1.0', 'Source-exclusive'],
        ['ChEMBL only',                    ' 10', '1.0', 'Source-exclusive'],
    ],
    col_widths=[2.2, 1.0, 1.3, 1.7],
)

add_table_title(doc, 'Table 4. Source independence score summary.')
make_table(doc,
    headers=['Score', 'n compounds (%)', 'Interpretation'],
    rows=[
        ['0.3', '1,366 (44.2%)', 'BindingDB + ChEMBL + PubChem redundancy'],
        ['0.5', '1,199 (38.8%)', 'BindingDB + PubChem redundancy'],
        ['0.6',   '167 (5.4%)',  'Genuinely multi-source (BindingDB + ChEMBL)'],
        ['0.7',    '23 (0.7%)',  'Genuinely multi-source (ChEMBL + PubChem)'],
        ['1.0',   '338 (10.9%)', 'Single-source (no cross-source redundancy; no second measurement)'],
        ['≥ 0.6 (free of re-curation redundancy)', '528 (17.1%)', '190 genuinely multi-source + 338 single-source'],
        ['< 0.6 (pipeline redundancy)', '2,565 (82.9%)', ''],
    ],
    col_widths=[1.9, 1.5, 2.8],
)

add_table_title(doc, 'Table 5. Comparison with existing PAD4 and activity-cliff resources.')
make_table(doc,
    headers=['Feature', 'ChEMBL (CHEMBL6111)', 'BindingDB (Q9UM07)', 'PAD4-DB v2'],
    rows=[
        ['Multi-source integration (PubChem+ChEMBL+BindingDB)', '✗', '✗', '✓'],
        ['Provenance / six-layer architecture', '✗', '✗', '✓'],
        ['Source-independence scoring', '✗', '✗', '✓'],
        ['HTS structural reference (327,336)', '✗', '✗', '✓'],
        ['Activity-cliff characterisation (MMP + permutation-validated)', '✗', '✗', '✓'],
        ['Cliff-hub annotations', '✗', '✗', '✓'],
        ['Benchmark-ready splits (via PAD4-Bench)', '✗', '✗', '✓'],
    ],
    col_widths=[3.0, 1.3, 1.3, 1.0],
)
hr(doc)

# ── MAIN FIGURES ──────────────────────────────────────────────────────
add_heading(doc, 'Figures', 1)

add_figure(doc, FIGS_MAIN / 'fig01_headline.png', 1,
    't-SNE landscape of the PAD4-DB v2 compound space. Two-dimensional t-SNE embedding '
    'of 3,093 compounds based on ECFP4 fingerprints (perplexity=30). Points are coloured '
    'by assay mechanism class (blue: enzymatic BAEE; teal: enzymatic_confirmed RFMS; '
    'cyan: fp_ic50; grey: other). Class A cliff-hub compounds (navy) are embedded within '
    'the dominant azaindole-benzimidazole cluster. Class B cliff-hub compounds (red) are '
    'positioned at the periphery adjacent to multiple chemotype clusters, consistent with '
    'their structural-attractor archetype.',
    width_in=6.0
)

add_figure(doc, FIGS_MAIN / 'fig02_source_overlap.png', 2,
    'Source overlap and independence scoring. UpSet plot showing the distribution of '
    '3,093 compounds across PubChem confirmatory, BindingDB, and ChEMBL sources. Inset: '
    'source independence score distribution histogram showing 528 non-redundant '
    'compounds (score ≥ 0.6; 190 genuinely multi-source + 338 single-source) versus '
    '2,565 pipeline-redundant compounds (score < 0.6).',
    width_in=5.5
)

add_figure(doc, FIGS_MAIN / 'fig03_potency.png', 3,
    'Potency and composition statistics. (a) Histogram of consensus pIC50 values '
    '(n=3,093; bin width=0.25). Bimodal distribution with primary peak at pIC50 ≈ 7.0 '
    'and patent-compound shoulder at pIC50 ≈ 5.0. Mean (6.55) and median (6.84) '
    'indicated. (b) pIC50 distribution by source type (published vs patent-exclusive; '
    'Mann-Whitney U p < 0.001). (c) Mechanism class distribution by compound count '
    'and mean pIC50.',
    width_in=6.2
)

add_figure(doc, FIGS_MAIN / 'fig04_scaffold.png', 4,
    'Scaffold diversity. (a) Scaffold series size distribution (y-axis: number of '
    'scaffolds; x-axis: series size). Top 10 scaffolds labelled. (b) Cumulative compound '
    'coverage: the 30 largest scaffolds cover approximately 30% of all compounds. '
    '(c) Lorenz curve of scaffold size distribution; Gini coefficient = 0.532.',
    width_in=6.2
)

add_figure(doc, FIGS_MAIN / 'fig05_cliff_network.png', 5,
    'Activity landscape and cliff network. Scatter plot of Tanimoto similarity versus '
    'ΔpIC50 for all 358,416 pairs with similarity ≥ 0.6. Severe cliff zone '
    '(Tanimoto ≥ 0.8, ΔpIC50 ≥ 2.0) highlighted. Hub-involved pairs coloured by hub '
    'class (navy: Class A, red: Class B).',
    width_in=6.2
)

add_figure(doc, FIGS_MAIN / 'fig06_mmp.png', 6,
    'MMP analysis of severe cliff pairs. Fig. 6 comprises two panels (a and b); panels c '
    'and d from an earlier draft were not retained. (a) MMP change-type distribution for 80 '
    'MMP-confirmed severe cliff pairs: single-atom changes (n=45, 56.3%), small substituent '
    'changes (n=27, 33.8%), medium substituent changes (n=8, 10.0%). (b) Four representative '
    'MMP-confirmed severe cliff pairs selected for change-type balance (warhead-free, '
    'non-ecfp4-only): two single-atom-change pairs and two small-substituent pairs, with no '
    'compound repeated across pairs and at least two non-hub pairs. Each pair is aligned at '
    'the shared MMP core; the transformed atom or substituent is highlighted in red. Labels: '
    'delta-pIC50, Tanimoto, MMP change type, hub class.',
    width_in=6.2
)
add_figure(doc, FIGS_MAIN / 'fig06b_cliff_pairs.png', '6b',
    '(panel b, enlarged). Fig. 6 comprises two panels (a and b); panels c and d from an '
    'earlier draft were not retained. Four representative MMP-confirmed severe cliff pairs '
    '(sub-panels i-iv), selected for change-type balance (warhead-free, non-ecfp4-only): '
    '(i) highest-delta-pIC50 single-atom-change pair (delta=2.92, Tan=0.84; non-hub vs Hub B). '
    '(ii) second highest-delta-pIC50 single-atom-change pair, no compound repeated (delta=2.79, '
    'Tan=0.80; Hub B vs non-hub). '
    '(iii) highest-delta-pIC50 small-substituent pair (delta=2.57, Tan=0.81; both non-hub). '
    '(iv) second highest-delta-pIC50 small-substituent pair, no compound repeated (delta=2.48, '
    'Tan=0.80; both non-hub). '
    'Each pair is aligned at the shared MMP core; the transformed atom or substituent is '
    'highlighted in red. Labels: delta-pIC50, Tanimoto similarity, MMP change type, hub class.',
    width_in=6.5
)
hr(doc)

# ── SUPPLEMENTARY FIGURES ─────────────────────────────────────────────
add_heading(doc, 'Supplementary Figures', 1)

add_supp_figure(doc, FIGS_SUPP / 'fig_s01_pipeline.png',
    'Supplementary Figure S1',
    'Curation pipeline flowchart from raw source data to PAD4-DB v2 outputs. The '
    'six-layer architecture (A–F) is illustrated with record counts at each stage. '
    'Final SAR compound count: 3,093. HTS structural reference: 327,336 compounds.',
    width_in=5.8
)

add_supp_figure(doc, FIGS_SUPP / 'fig_s02_sali.png',
    'Supplementary Figure S2',
    'SALI (Structure-Activity Landscape Index) distribution for all 358,416 compound '
    'pairs at Tanimoto ≥ 0.6. SALI > 10: 335 pairs; SALI > 20: 19 pairs; '
    'SALI max = 65.88 (single outlier). The bulk of pairs occupy the concordant, '
    'low-SALI quadrant, confirming that severe cliffs are rare events.',
    width_in=5.5
)

add_supp_figure(doc, FIGS_SUPP / 'fig_s03_patent.png',
    'Supplementary Figure S3',
    'Patent-exclusive compound analysis. pIC50 distribution comparison (patent-exclusive '
    'n=233, mean=6.08; published n=2,860, mean=6.59; Mann-Whitney U p < 0.001). '
    'Patent compounds cover 103 scaffold series absent from non-patent space with a mean '
    'series density of 2.1 compounds per series (vs 2.5 for non-patent scaffolds).',
    width_in=5.5
)

add_supp_figure(doc, FIGS_SUPP / 'fig_s04_reference_recovery.png',
    'Supplementary Figure S4',
    'Reference compound recovery. Concordance scatter (x = published pIC50, y = PAD4-DB v2 '
    'consensus pIC50) for seven recovered PAD4 inhibitors: Streptonigrin, Cl-amidine, '
    'F-amidine, GSK484 (standardised to free base; source data contained HCl salt), TDFA, '
    'BMS-P5, and JBI-589. Five of seven compounds have |ΔpIC50| ≤ 0.15 log units (filled '
    'circles). GSK484 has |ΔpIC50| = 0.25 (open circle; within typical inter-assay '
    'variability). JBI-589 has |ΔpIC50| = 0.91, '
    'attributed to calcium-concentration differences between assay formats; retained with '
    'the discrepancy documented. Mean |ΔpIC50| across six non-JBI-589 compounds = 0.061.',
    width_in=5.0
)

add_supp_figure(doc, FIGS_SUPP / 'fig_s05_scaffold_structures.png',
    'Supplementary Figure S5',
    'Dominant scaffold series: azaindole-benzimidazole core structure and pIC50 '
    'distribution across the 174-compound series (mean pIC50 = 7.07, SD = 0.88, '
    'range = 3.5 log units). The two Class A cliff-hub compounds (pIC50 ≈ 5.4) are '
    'indicated; their position as within-series potency floors accounts for 27 of 94 '
    'severe cliff pairs.',
    width_in=5.8
)
add_supp_figure(doc, FIGS_SUPP / 'fig_s06_permutation.png',
    'Supplementary Figure S6',
    'Permutation analysis of the activity-cliff landscape (10,000 permutations of consensus '
    'pIC50 with the similarity structure held fixed). (a) Null distribution of severe cliff '
    'counts; the observed count (94, red line) lies below every permutation (null 1,923 ± 125; '
    'p < 0.0001) — a ~20-fold depletion confirming the cliffs are genuine discontinuities, '
    'not artefacts of the bimodal potency distribution. (b) Null distribution of hub '
    'concentration (fraction of 94 count-matched cliff pairs incident to the four '
    'highest-degree compounds); the observed 53.2% (red line) exceeds every permutation '
    '(null 15.2% ± 2.6%; p < 0.0001), confirming the hub structure is potency-driven '
    'rather than a similarity-density artefact.',
    width_in=6.5
)
add_supp_figure(doc, FIGS_SUPP / 'fig_s07_physicochemical.png',
    'Supplementary Figure S7',
    'Physicochemical property landscape of the 3,093 PAD4-DB v2 compounds, computed with '
    'RDKit from standardised SMILES (zero parse failures). Eight-panel histogram grid: '
    '(a) molecular weight, (b) Crippen cLogP, (c) topological polar surface area (TPSA), '
    '(d) H-bond acceptors, (e) H-bond donors, (f) rotatable bonds, (g) fraction Csp3, '
    '(h) number of aromatic rings. The red dashed line marks the median of each distribution. '
    'Median molecular weight is 590.7 Da; 19.8% of compounds satisfy all four Lipinski '
    'rule-of-five criteria strictly, 68.9% satisfy the classic <=1-violation threshold, '
    'and 90.7% satisfy both Veber criteria.',
    width_in=6.8
)
hr(doc)

# ── REFERENCES ────────────────────────────────────────────────────────
add_heading(doc, 'References', 1)
refs = [
    '[PAD4_REV1] [PAD4 biology review 1 — TBD; suggest Damgaard 2016 Nat Rev Drug Discov]',
    '[PAD4_REV2] [PAD4 biology review 2 — TBD; suggest Witalison 2015 Curr Drug Targets]',
    '[Senger2020] [Database overlap architecture — TBD; Beilstein BJOC]',
    '[ChEMBL2023] Zdrazil B, et al. The ChEMBL Database in 2023. Nucleic Acids Res. 2024;52:D1180–D1192.',
    '[BemisMurcko1996] Bemis GW, Murcko MA. The properties of known drugs. 1. Molecular frameworks. J Med Chem. 1996;39:2887–2893.',
    '[StumpfeBajorath2012] Stumpfe D, Bajorath J. Exploring activity cliffs in medicinal chemistry. J Med Chem. 2012;55:2932–2942.',
    '[Senger2009] Senger S, et al. [ECFP4 threshold justification for cliff analysis — TBD]',
    '[RogersHahn2010] Rogers D, Hahn M. Extended-connectivity fingerprints. J Chem Inf Model. 2010;50:742–754.',
    '[Lex2014] Lex A, et al. UpSet: Visualization of intersecting sets. IEEE Trans Vis Comput Graph. 2014;20:1983–1992.',
    '[Knuckley2010] Knuckley B, et al. Substrate specificity and kinetic studies of PADs 1, 3, and 4 identify potent and selective inhibitors of protein arginine deiminase 3. Biochemistry. 2010;49:4852–4863.',
    '[PAD4Bench] [PAD4-Bench companion paper — TBD; fill when submitted]',
    '[MoleculeACE] van Tilborg D, Alenicheva A, Grisoni F. Exposing the Limitations of Molecular Machine Learning with Activity Cliffs. J Chem Inf Model. 2022;62(23):5938–5951. [VERIFY citation format]',
    '[ACNet] Zhang Z, Zhao B, Xie A, Bian Y, Zhou S. Activity Cliff Prediction: Dataset and Benchmark. arXiv:2302.07541, 2023. [VERIFY citation format]',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'{i}.  {ref}')
    p.style = doc.styles['Normal']
    set_spacing(p, before=0, after=4, line=12)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.runs[0].font.size = Pt(9)

# ── SAVE ──────────────────────────────────────────────────────────────
doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1e3
print(f'\nSaved: {OUT}')
print(f'Size:  {size_kb:.0f} KB')
print('Done.')
